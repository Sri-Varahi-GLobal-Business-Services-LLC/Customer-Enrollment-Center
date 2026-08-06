# -*- coding: utf-8 -*-
"""Append Appendix A (UI Technology Options Considered) to Solution Design v2.0."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = r"C:\Users\jnamm\OneDrive\Desktop\Service Cloud for Utilities\Enrollment Center\Enrollment Center - SAP Service Cloud V2 Solution Design v2.0.docx"
HEADER_FILL="1F4E79"; ALT_FILL="DEEAF6"
doc = Document(PATH)

doc.add_page_break()
h = doc.add_heading("Appendix A — UI Technology Options Considered", level=1)

def para(text, bold=False):
    pr = doc.add_paragraph(); r = pr.add_run(text); r.bold = bold
    r.font.name = "Arial"; r.font.size = Pt(10.5)
    pr.paragraph_format.space_after = Pt(6)
    return pr

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def table(headers, rows, widths, font_size=8.5):
    t = doc.add_table(rows=1, cols=len(headers)); t.style="Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,hd in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""
        r=c.paragraphs[0].add_run(hd); r.bold=True; r.font.size=Pt(font_size)
        r.font.name="Arial"; r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); shade(c,HEADER_FILL)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=""
            run=cells[i].paragraphs[0].add_run(str(val)); run.font.size=Pt(font_size); run.font.name="Arial"
            if ri%2==1: shade(cells[i],ALT_FILL)
    for i,w in enumerate(widths):
        for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

para("Three UI technology options were evaluated for the Enrollment Center application. Note that 'Fiori' is a "
     "design system, not a hosting decision: the recommended option delivers a Fiori-design application "
     "(SAPUI5, Horizon theme) hosted on SAP BTP. The full weighted evaluation is documented in "
     "'Enrollment Center – UI Technology Evaluation v1.0'.")
table(["Option","Description","Assessment"], [
 ["A. SAP BTP side-by-side (RECOMMENDED)",
  "SAPUI5/Fiori Elements (TypeScript, Horizon) on HTML5 Application Repository; CAP services on Cloud Foundry; embedded in SC V2 as work-center mashup and Move-In step",
  "Selected. Clean core; native SC V2 embedding with IAS SSO; independent release cycle; the orchestration (eventing, workflow, deferred fulfillment, audit) requires BTP services regardless; API layer reusable for self-service/IVR/DER channels. Carries BTP subscription cost (Config Guide Section 2.2) — justified by the SC V2 strategy."],
 ["B. Embedded Fiori (RAP) on S/4HANA",
  "Fiori/UI5 app developed on the S/4HANA ABAP stack (RAP), surfaced to agents from the S/4 gateway",
  "Fallback only (client without BTP). Agents work in SC V2, so browsers would need routes to the S/4HANA gateway; custom UI enters the digital core; UI releases coupled to S/4HANA maintenance; BTP still needed for workflow/eventing halves; no channel reuse. Reduced-scope variant documented in the evaluation."],
 ["C. ABAP Web Dynpro",
  "Classic Web Dynpro ABAP application in S/4HANA",
  "Rejected. Maintenance-mode technology outside SAP's strategic direction; no SC V2 embedding or Horizon theming; desktop-only UX; custom code in the core; unsuitable for any new build."],
], widths=[1.4,2.4,2.7])
para("Decision criteria (weighted in the evaluation document): SC V2 embedding quality, clean-core compliance, "
     "fit for the orchestration architecture, channel reuse, release independence, SAP strategic alignment, "
     "skills availability, total cost, performance and security. Option A scores highest on all criteria "
     "except subscription cost. This appendix records the decision for architecture review; the detailed "
     "scoring, cost discussion and risk register are maintained in the evaluation document.")

doc.save(PATH)
print("Appendix A appended to SDD")
