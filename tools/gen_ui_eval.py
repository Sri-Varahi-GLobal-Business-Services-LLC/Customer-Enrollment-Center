# -*- coding: utf-8 -*-
"""Generate Enrollment Center UI Technology Evaluation document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT="1F4E79"; HEADER_FILL="1F4E79"; ALT_FILL="DEEAF6"
doc = Document()
n = doc.styles["Normal"]; n.font.name="Arial"; n.font.size=Pt(10.5); n.paragraph_format.space_after=Pt(6)
for lvl, sz in [("Heading 1",16),("Heading 2",13),("Heading 3",11.5)]:
    st = doc.styles[lvl]; st.font.name="Arial"; st.font.size=Pt(sz); st.font.bold=True
    st.font.color.rgb = RGBColor(0x1F,0x4E,0x79)
    rpr = st.element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
    if rf is None: rf = OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a),'Arial')

def p(text, bold=False, size=None, color=None, align=None):
    para = doc.add_paragraph(); r = para.add_run(text); r.bold=bold
    if size: r.font.size=Pt(size)
    if color: r.font.color.rgb=RGBColor.from_string(color)
    if align: para.alignment=align
    return para

def bullets(items):
    for it in items:
        if isinstance(it, tuple):
            para = doc.add_paragraph(style="List Bullet")
            r = para.add_run(it[0]+" — "); r.bold=True
            para.add_run(it[1])
        else:
            doc.add_paragraph(it, style="List Bullet")

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def table(headers, rows, widths=None, font_size=8.5):
    t = doc.add_table(rows=1, cols=len(headers)); t.style="Table Grid"
    t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""
        r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(font_size)
        r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); shade(c,HEADER_FILL)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=""
            run=cells[i].paragraphs[0].add_run(str(val)); run.font.size=Pt(font_size)
            if ri%2==1: shade(cells[i],ALT_FILL)
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

# ===== COVER =====
for _ in range(6): doc.add_paragraph()
p("Customer Enrollment Center", bold=True, size=26, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER)
p("UI Technology Evaluation", bold=True, size=18, color="404040", align=WD_ALIGN_PARAGRAPH.CENTER)
p("SAP BTP side-by-side vs. embedded Fiori (RAP) on S/4HANA vs. ABAP Web Dynpro — weighted evaluation, cost, risks and decision",
  size=12, color="595959", align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(7): doc.add_paragraph()
table(["Attribute","Value"], [
 ["Document","UI Technology Evaluation (referenced by Solution Design v2.0, Appendix A)"],
 ["Version / Status","1.0 (Draft for Review)"],
 ["Date","08 July 2026"],
 ["Decision","Option A — SAP BTP side-by-side extension with Fiori-design SAPUI5 UI"],
 ["Audience","Architecture review board, IT leadership, program management"],
], widths=[1.7,4.8], font_size=9)
doc.add_page_break()

# ===== 1 =====
doc.add_heading("1. Purpose and Framing", level=1)
p("This document records the evaluation of UI/application technology options for the Customer Enrollment "
  "Center and the rationale for the selected option, for architecture governance. One framing point removes "
  "most confusion in this discussion: 'SAP Fiori' is SAP's design system and UX standard, not a hosting "
  "location. A Fiori application can run on the S/4HANA ABAP stack or on SAP BTP. The recommended option IS a "
  "Fiori application — SAPUI5 with the Horizon theme — hosted on BTP. The evaluation therefore compares "
  "application platforms, not visual designs: all compliant options would look like Fiori.")
p("Requirements recap that drives the evaluation (from Solution Design v2.0):", bold=True)
bullets([
 "Embedded in SAP Service Cloud V2 at two access points (work center mashup + Move-In in-screen step) with silent SSO.",
 "Heavy orchestration beyond the UI: single-call eligibility, deferred move-in fulfillment via events, supervisor workflow, audit, exception handling, DER-platform handoff.",
 "Clean core: no UI custom code inside S/4HANA; S/4HANA extended only via released APIs and governed wrappers.",
 "Channel-agnostic service layer (EC-01…EC-09) reusable for future self-service, IVR and DER platform.",
 "Program catalog changes at business speed, decoupled from S/4HANA release calendar.",
])

# ===== 2 OPTIONS =====
doc.add_heading("2. Options Described", level=1)
table(["Option","Stack","Where It Runs","How Agents Reach It"], [
 ["A. BTP side-by-side (recommended)","SAPUI5/Fiori Elements (TypeScript, Horizon) + CAP (Node.js) + HANA Cloud + Integration Suite + Event Mesh + SBPA","SAP BTP Cloud Foundry / HTML5 App Repository","SC V2 mashup / embedded step; IAS SSO; public cloud URL"],
 ["B. Embedded Fiori (RAP) on S/4HANA","ABAP RESTful Application Programming (RAP) + Fiori Elements on the S/4HANA gateway","S/4HANA application server (custom Z code in the core landscape)","iFrame/mashup pointing at the S/4 gateway — requires network path from agent browsers to S/4HANA"],
 ["C. ABAP Web Dynpro","Web Dynpro ABAP + FPM","S/4HANA application server","Direct URL to ITS/WD service on S/4HANA"],
 ["A2 (variant). SAP Build Apps on BTP","Low-code UI on BTP consuming the same CAP APIs","SAP BTP","Same as A; considered as a UI-layer variant of Option A, not a separate platform"],
], widths=[1.5,2.3,1.5,1.7], font_size=8)

# ===== 3 CRITERIA =====
doc.add_heading("3. Evaluation Criteria and Weights", level=1)
table(["Criterion","Weight","Why It Matters Here"], [
 ["SC V2 embedding & SSO quality","20%","The agent desktop IS SC V2; both access points must feel native (Horizon, silent SSO, context handover)"],
 ["Fit for orchestration architecture","20%","Eligibility facade, eventing, workflow, deferred fulfillment, audit — the majority of the build"],
 ["Clean core / upgrade safety","15%","Client policy; S/4HANA upgrade cadence must not be constrained by this app"],
 ["Channel reuse of services","10%","Self-service / IVR / DER platform roadmap"],
 ["Release independence","10%","Seasonal program changes vs. S/4HANA maintenance windows"],
 ["SAP strategic alignment / longevity","10%","10-year investment protection"],
 ["Skills & delivery speed","5%","Available team skills (UI5/CAP vs ABAP)"],
 ["Run cost (subscriptions/licensing)","5%","BTP service subscriptions vs. embedded 'free' hosting"],
 ["Performance & security topology","5%","Latency to data; exposure of backend endpoints to browsers"],
], widths=[2.3,0.8,3.4])

# ===== 4 SCORING =====
doc.add_heading("4. Scoring Matrix (1 = poor, 5 = excellent)", level=1)
table(["Criterion (weight)","A. BTP","B. RAP on S/4","C. Web Dynpro"], [
 ["SC V2 embedding & SSO (20%)","5","2","1"],
 ["Orchestration fit (20%)","5","2","1"],
 ["Clean core / upgrade safety (15%)","5","2","1"],
 ["Channel reuse (10%)","5","2","1"],
 ["Release independence (10%)","5","2","2"],
 ["SAP strategic alignment (10%)","5","4","1"],
 ["Skills & delivery speed (5%)","4","4","3"],
 ["Run cost (5%)","3","4","4"],
 ["Performance & security topology (5%)","4","3","2"],
 ["WEIGHTED TOTAL","4.80","2.45","1.30"],
], widths=[2.9,1.2,1.3,1.3], font_size=9)
p("Sensitivity: Option B only approaches Option A if the weights for embedding, orchestration and channel "
  "reuse are near zero — i.e., a different project. No plausible weighting rescues Option C.")

# ===== 5 ANALYSIS =====
doc.add_heading("5. Analysis Detail per Option", level=1)
doc.add_heading("5.1 Option A — BTP Side-by-Side (Recommended)", level=2)
bullets([
 ("Embedding","cloud URL embeds natively in SC V2 mashups; IAS gives silent SSO in the agent session; Horizon theme matches the surrounding desktop."),
 ("Orchestration","CAP + Event Mesh + SBPA + HANA Cloud are exactly the services the A1–A11 runtime flows require; the UI is the thin top of an architecture that must be on BTP anyway."),
 ("Clean core","S/4HANA receives only governed wrappers, BRFplus content and released-API consumption (TDD D-01…D-20); upgrades unaffected."),
 ("Costs","subscription components per Config Guide 2.2; partially shared with other CX extensions on the same subaccounts; the eligibility/enrollment API layer amortizes across future channels."),
 ("Risks & mitigations","platform skills (mitigate: CAP/UI5 enablement, CI/CD templates); subscription creep (mitigate: right-sized plans, quarterly consumption review); SC V2 embedding contract (mitigate: Discover-phase spike, already planned)."),
])
doc.add_heading("5.2 Option B — Embedded Fiori (RAP) on S/4HANA", level=2)
bullets([
 ("Where it wins","no new platform if the client truly has zero BTP; direct data locality (no integration hop for reads); ABAP team familiarity."),
 ("Where it fails here","agents' browsers need routed, authenticated access to the S/4HANA gateway (network exposure + SSO complexity from SC V2 context); custom app code lands in the core landscape; every UI change rides S/4HANA transports and maintenance windows; the workflow/eventing/deferred-fulfillment half of the solution still needs BTP or heavy ABAP custom equivalents (BPEM/workflow), eroding the 'no new platform' argument; enrollment APIs are not reusable by cloud channels without adding the integration layer anyway."),
 ("When to choose it","documented fallback: client rejects BTP subscriptions entirely AND accepts reduced scope (no event-driven deferred fulfillment, workflow via S/4 inbox, no channel reuse). Recorded as a conscious scope trade, not an equivalent."),
])
doc.add_heading("5.3 Option C — ABAP Web Dynpro", level=2)
bullets([
 "Maintenance-mode technology: excluded from SAP's strategic UX direction (Fiori/Horizon); not enhanced for years; skills pool shrinking.",
 "No SC V2 embedding story: no Horizon theme, poor iframe behavior, desktop-only rendering, no mobile support.",
 "All Option B disadvantages (core custom code, release coupling, network exposure) plus an obsolete UI framework on top.",
 "Verdict: rejected outright; would also signal outdated architecture to any review board.",
])
doc.add_heading("5.4 Variant A2 — SAP Build Apps (low-code UI on BTP)", level=2)
p("Same platform decision as Option A; only the UI layer differs. Assessment: viable for a first release if "
  "the client wants citizen-developer ownership, but the four-panel table density, wizard flows and embedded "
  "move-in mode favor pro-code SAPUI5/Fiori Elements. The CAP API layer is identical either way, so this "
  "choice is reversible and deferred to the Explore phase. Default: SAPUI5.")

# ===== 6 COST =====
doc.add_heading("6. Cost Considerations (qualitative)", level=1)
table(["Cost Element","A. BTP","B. RAP on S/4"], [
 ["Platform subscriptions","BTP services per Config Guide 2.2 (CF runtime, HANA Cloud, Integration Suite, Event Mesh, SBPA, misc.) — partially shared with other extensions","None new — but Integration Suite/Event Mesh still required for the event flows, or scope is cut"],
 ["Build effort","CAP+UI5 build; wrappers in S/4 (D-01…D-20)","Similar S/4 development volume PLUS RAP UI app; workflow/eventing replacements add effort"],
 ["Run / upgrade effort","Independent pipeline; zero S/4 upgrade impact","Regression testing of custom UI at every S/4 SP/upgrade"],
 ["Channel reuse economics","API layer amortized over future channels","Re-build required for each new channel"],
], widths=[1.7,2.5,2.4], font_size=8.5)
p("Note: exact subscription pricing depends on the client's BTP commercial model (CPEA/BTPEA credits vs. "
  "subscription) and is provided by the account team; this evaluation deliberately stays qualitative.")

# ===== 7 DECISION =====
doc.add_heading("7. Decision and Governance Record", level=1)
bullets([
 ("Decision","Option A — SAP BTP side-by-side extension; UI as Fiori-design SAPUI5 application (variant A2 Build Apps reversible at Explore phase)."),
 ("Conditions","BTP subaccounts per Config Guide Section 2; Discover-phase spike validates both SC V2 embedding points (already planned in Solution Design Section 15)."),
 ("Fallback trigger","only if BTP adoption is rejected at executive level: Option B with the documented scope reductions — requires re-approval of the Solution Design."),
 ("Explicitly rejected","Option C (Web Dynpro) — not to be revisited."),
 ("Document impact","none — Solution Design v2.0, TDD v1.1, Config Guide v1.0, BRFplus v1.1 and Runtime Interactions v1.0 all assume Option A."),
])

out = r"C:\Users\jnamm\OneDrive\Desktop\Service Cloud for Utilities\Enrollment Center\Enrollment Center - UI Technology Evaluation v1.0.docx"
doc.save(out)
print("Saved:", out)
