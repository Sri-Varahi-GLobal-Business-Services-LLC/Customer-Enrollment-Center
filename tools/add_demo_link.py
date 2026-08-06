# -*- coding: utf-8 -*-
"""Add interactive demo link to Enrollment Center SDD v2.0 (cover table + UI section)."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = r"C:\Users\jnamm\OneDrive\Desktop\Service Cloud for Utilities\Enrollment Center\Enrollment Center - SAP Service Cloud V2 Solution Design v2.0.docx"
URL = "https://claude.ai/code/artifact/d4ef964a-7610-4bd9-b581-921b1a2a7a3f"

doc = Document(PATH)

# ---- 1. Add row to cover attribute table (first table) ----
cover = doc.tables[0]
row = cover.add_row()
labels = ["UI Prototype",
          "Interactive UI demo (illustrative prototype only, not a specification): " + URL]
for i, txt in enumerate(labels):
    cell = row.cells[i]
    cell.text = ""
    run = cell.paragraphs[0].add_run(txt)
    run.font.size = Pt(9.5)
    run.font.name = "Arial"
# match column widths of existing rows
for i in range(2):
    row.cells[i].width = cover.rows[0].cells[i].width

# ---- 2. Insert prototype note in section 4.4 before the 'Technology choice:' paragraph ----
target = None
for para in doc.paragraphs:
    if para.text.startswith("Technology choice:"):
        target = para
        break
assert target is not None, "Technology choice paragraph not found"

new_p = OxmlElement("w:p")
target._p.addprevious(new_p)
from docx.text.paragraph import Paragraph
np = Paragraph(new_p, target._parent)
run = np.add_run(
    "An interactive prototype of this UI is available for stakeholder walkthroughs at " + URL +
    " . It demonstrates both access points (work center and Move-In embedded mode), the four panels, "
    "the status model including deferred fulfillment on move-in activation and cancellation cascade, "
    "eligibility message severities, supervisor override for warning-level failures, and de-enrollment "
    "exit-rule validation. The prototype is illustrative only; this document and the program-specific "
    "functional specifications remain the binding design."
)
run.font.name = "Arial"
run.font.size = Pt(10.5)
np.paragraph_format.space_after = Pt(6)

doc.save(PATH)
print("Updated:", PATH)
