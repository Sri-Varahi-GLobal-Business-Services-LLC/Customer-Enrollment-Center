# -*- coding: utf-8 -*-
"""Update SDD v2.0 enrollment scope with US reference catalog programs and DPP/DPA distinction."""
from docx import Document

PATH = r"C:\Users\jnamm\OneDrive\Desktop\Service Cloud for Utilities\Enrollment Center\Enrollment Center - SAP Service Cloud V2 Solution Design v2.0.docx"
doc = Document(PATH)

REPL = [
 ("Peak-Time Rebate (PTR), Critical Peak Pricing (CPP), time-of-use rates",
  "Peak-Time Rebate (PTR), Critical Peak Pricing (CPP), Dynamic Peak Pricing (DPP), time-of-use rates, EV managed charging, green power / community solar options"),
 ("Deferred Payment Plan (DPP), budget billing, PrePay",
  "Deferred Payment Arrangement (DPA), budget billing, PrePay, AutoPay, payment extension, arrearage management (AMP)"),
 ("Low-income assistance (DLA-type), medical alert",
  "Low-income assistance (CARE/FERA, LIHEAP, PIPP), medical alert, third-party bill notification"),
 ("Customer Selects Due Date (CSDD), eBill / paperless billing",
  "Customer Selects Due Date (CSDD), eBill / paperless billing, summary/consolidated billing, usage and high-bill alerts"),
 ("proposed: DPP, PTR, CPP, CSDD, PrePay based on legacy scope",
  "proposed: DPP (Dynamic Peak Pricing), DPA, PTR, CPP, CSDD, PrePay based on legacy scope; extended catalog per the U.S. Utility Customer Programs Reference (Aug 2026). Note: in the legacy source, 'DPP' denotes Dynamic Peak Pricing (a rate program fulfilled by product change); the installment payment plan is designated DPA"),
]

count = 0
def fix(para):
    global count
    for run in para.runs:
        for old, new in REPL:
            if old in run.text:
                run.text = run.text.replace(old, new); count += 1

for p in doc.paragraphs: fix(p)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            for p in c.paragraphs: fix(p)

doc.save(PATH)
print("Replacements:", count)
