# -*- coding: utf-8 -*-
"""Generate Enrollment Center BRFplus Application Detailed Configuration document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT="1F4E79"; HEADER_FILL="1F4E79"; ALT_FILL="DEEAF6"
doc = Document()
n = doc.styles["Normal"]; n.font.name="Arial"; n.font.size=Pt(10.5); n.paragraph_format.space_after=Pt(6)
for lvl, sz in [("Heading 1",16),("Heading 2",13),("Heading 3",11.5)]:
    st = doc.styles[lvl]; st.font.name="Arial"; st.font.size=Pt(sz); st.font.bold=True
    st.font.color.rgb = RGBColor(0x1F,0x4E,0x79)
    rpr = st.element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
    if rf is None: rf = OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a),'Arial')

def p(text, bold=False, size=None, color=None, align=None):
    para = doc.add_paragraph(); r = para.add_run(text); r.bold=bold
    if size: r.font.size=Pt(size)
    if color: r.font.color.rgb=RGBColor.from_string(color)
    if align: para.alignment=align
    return para

def mono(text):
    para = doc.add_paragraph(); r = para.add_run(text)
    r.font.name="Consolas"; r.font.size=Pt(8.5)
    para.paragraph_format.space_after=Pt(1)
    return para

def bullets(items):
    for it in items:
        if isinstance(it, tuple):
            para = doc.add_paragraph(style="List Bullet")
            r = para.add_run(it[0]+" — "); r.bold=True
            para.add_run(it[1])
        else:
            doc.add_paragraph(it, style="List Bullet")

def numbered(items):
    for it in items: doc.add_paragraph(it, style="List Number")

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def table(headers, rows, widths=None, font_size=8.5):
    t = doc.add_table(rows=1, cols=len(headers)); t.style="Table Grid"
    t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""
        r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(font_size)
        r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); shade(c,HEADER_FILL)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=""
            run=cells[i].paragraphs[0].add_run(str(val)); run.font.size=Pt(font_size)
            if ri%2==1: shade(cells[i],ALT_FILL)
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

# ============ COVER ============
for _ in range(6): doc.add_paragraph()
p("Enrollment Center — BRFplus Application", bold=True, size=26, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER)
p("Detailed Configuration Specification", bold=True, size=18, color="404040", align=WD_ALIGN_PARAGRAPH.CENTER)
p("Application ZEC_PROGRAM_ELIGIBILITY — data objects, expressions, decision tables, rulesets, messages, invocation, testing and transport",
  size=12, color="595959", align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(7): doc.add_paragraph()
table(["Attribute","Value"], [
 ["Document","BRFplus Detailed Configuration (companion to TDD v1.1 Section 6 / Config Guide D-16)"],
 ["Version / Status","1.1 (adds Section 11 object inventory; consolidated decision-table model; companion configuration workbook)"],
 ["Date","08 July 2026"],
 ["Legacy Migration","FUNC_PGMENRL_PRECHECK invocation pattern and ruleset RS_DPP_INT_VALIDATION (Dynamic Peak Pricing) migrated into this application"],
 ["Audience","ABAP/BRFplus developers, business rule owners, QA"],
], widths=[1.7,4.8], font_size=9)
doc.add_page_break()

# ============ 1 ============
doc.add_heading("1. Application Setup", level=1)
table(["Attribute","Value"], [
 ["Application name","ZEC_PROGRAM_ELIGIBILITY"],
 ["Description","Enrollment Center — Program Eligibility and Exit Rules"],
 ["Storage type","S (System) — transportable, client-independent rule structure"],
 ["Development package","ZEC_ENROLLMENT"],
 ["Software component / transport","Workbench request; 'Create Local Application' = No"],
 ["Application component","IS-U-CS (customer service)"],
 ["Language","EN master; texts translatable"],
 ["Application exit class","ZCL_EC_FDT_EXIT (element texts for message rendering; authorization check hook)"],
], widths=[2.0,4.5], font_size=9)
p("Creation: transaction BRFplus (or BRF+) → Workbench → Create Application with the values above. All objects "
  "below are created inside this application. Naming: functions F_*, data objects CTX_*/RES_*/EL_*, "
  "expressions EX_*, decision tables DT_*, rulesets RS_*, rules RU_*.")

# ============ 2 DATA OBJECTS ============
doc.add_heading("2. Data Objects", level=1)
doc.add_heading("2.1 Context Structure CTX_EC_ELIG", level=2)
p("Top-level context passed by Z_EC_ELIGIBILITY_PRECHECK. DDIC-bound where a dictionary type exists.")
table(["Element","Type (DDIC)","Description"], [
 ["BUSINESS_PARTNER","BU_PARTNER","Business partner"],
 ["MODE","CHAR1 (S=Standalone, P=Pre-Active)","Evaluation mode; pre-active downgrades history-dependent hard stops to 'not yet eligible'"],
 ["MOVEIN_ID","Move-in document ID","Filled in mode P"],
 ["PROGRAM_FILTER","Table of ZEC_PROGRAM_ID","Optional; empty = all active catalog programs"],
 ["T_CA","Table of CTX_EC_CA (2.2)","One row per contract account in scope"],
 ["T_HISTORY","Table of CTX_EC_HIST (2.3)","Program participation history"],
 ["T_ACTIVE","Table of ZEC_PROGRAM_ID + CA","Currently active participations (for exclusion rules)"],
 ["REQUEST_DATE","DATS","Evaluation date (sy-datum default)"],
], widths=[1.7,1.7,3.1])
doc.add_heading("2.2 Per-Contract-Account Structure CTX_EC_CA", level=2)
table(["Element","Type","Filled From (by the RFC wrapper)"], [
 ["CONTRACT_ACCOUNT","VKONT_KK","FKKVKP"],
 ["DIVISION","SPARTE_D","EVER"],
 ["RATE_CATEGORY","TARIFTYP","EANLH (current)"],
 ["RATE_CLASS","CHAR3 (RES/COM)","Derived via EX_RATE_CLASS"],
 ["CONTRACT_ACTIVE","BOOLE_D","EVER (mode P: false)"],
 ["AMI_METER","BOOLE_D","I_UtilsAdvncdMeteringSystem / installed device"],
 ["AMI_COMM_OK","BOOLE_D","Device comm status (meter ticket check)"],
 ["DUNNING_LEVEL","MAHNS_D","FI-CA highest open dunning level"],
 ["RETURNED_PAYMENTS_12M","INT2","FI-CA returns history count"],
 ["ARREARS_AMOUNT","BETRW_KK","FI-CA overdue open items sum"],
 ["ACTIVE_INSTPLAN","BOOLE_D","FI-CA installment plan check"],
 ["BROKEN_PLANS_24M","INT2","Deactivated-for-default plans count"],
 ["BUDGET_BILLING","BOOLE_D","EABP active plan"],
 ["PREPAY_ACTIVE","BOOLE_D","Prepay settings on CA"],
 ["MEDICAL_FLAG","BOOLE_D","Protected-customer/life-support mark (BP/CA)"],
 ["EBILL_ACTIVE","BOOLE_D","Correspondence settings"],
 ["BILLING_HISTORY_MONTHS","INT2","Months of billing docs at premise (ERCH)"],
 ["PAYMENT_EXT_ACTIVE","BOOLE_D","Open due-date deferral"],
 ["PAYMENT_EXT_12M","INT2","Extensions granted last 12 months"],
 ["CURTAILABLE_LOAD_KW","INT4","C&I: declared/measured curtailable load"],
 ["CA_COUNT_BP","INT2","Number of CAs of the BP (summary billing rule)"],
], widths=[1.9,1.3,3.3], font_size=8)
doc.add_heading("2.3 History Structure CTX_EC_HIST and Result Table RES_EC_ELIG", level=2)
table(["Structure","Elements"], [
 ["CTX_EC_HIST","PROGRAM_ID; CONTRACT_ACCOUNT; START_DATE; END_DATE; OUTCOME (C=Completed, D=De-enrolled, X=Cancelled, F=Default/Removed)"],
 ["RES_EC_ELIG (table row)","PROGRAM_ID; CONTRACT_ACCOUNT; ELIGIBLE (Y/N/I); SEVERITY (E=hard stop, W=warning, I=not yet eligible); MSGID (ZEC_ELIG); MSGNO; MSGV1–MSGV4; WAITING_END (DATS, when applicable)"],
], widths=[1.7,4.8], font_size=9)
p("Result contract: one row per failed/limited check (a program may have several rows); programs with no rows "
  "for a CA are eligible on that CA. All rulesets append to the same result table — no first-fail exit, per "
  "the legacy design.")

# ============ 3 FUNCTIONS ============
doc.add_heading("3. Functions", level=1)
table(["Function","Mode","Signature","Processing"], [
 ["F_EC_ELIGIBILITY_ALL","Event Mode","Context: CTX_EC_ELIG; Result: RES_EC_ELIG (table)","Triggers all assigned rulesets (Section 6); each ruleset has a precondition on PROGRAM_FILTER so only requested programs execute; rulesets loop over T_CA"],
 ["F_EC_DEENROLL_CHECK","Functional Mode","Context: PROGRAM_ID, CONTRACT_ACCOUNT, enrollment dates, CTX_EC_CA row; Result: RES_EC_ELIG","Top expression EX_EXIT_RULES (DT_EXIT_RULES + FI-CA balance lookup)"],
], widths=[1.7,1.0,1.9,1.9], font_size=8.5)

# ============ 4 EXPRESSIONS ============
doc.add_heading("4. Reusable Expressions", level=1)
table(["Expression","Type","Input → Output","Logic"], [
 ["EX_RATE_CLASS","Decision table (DT_RATECLASS_MATRIX)","RATE_CATEGORY → RATE_CLASS","Pattern match rate category to RES/COM (maintained matrix, Section 5.1)"],
 ["EX_WAITING_PERIOD","Loop + DT_WAITING_PERIODS","PROGRAM_ID, T_HISTORY → WAITING_END","Latest relevant history row; adds waiting months per program/outcome; returns end date if in the future"],
 ["EX_PROGRAM_EXCLUSION","Table operation + DT_PROGRAM_EXCLUSIONS","PROGRAM_ID, T_ACTIVE → message keys","Mutual exclusion lookup (Section 5.3)"],
 ["EX_AMI_CHECK","Formula on CTX_EC_CA","AMI_METER, AMI_COMM_OK → flags","Pure context read (values pre-fetched by wrapper for performance)"],
 ["EX_HIST_MONTHS_OK","Formula","BILLING_HISTORY_MONTHS ≥ threshold parameter","Threshold per calling rule (12 for BB)"],
 ["EX_SEVERITY_MODE","Formula","MODE, rule severity → effective severity","In mode P: rules flagged 'history/contract-dependent' return I instead of E (implements pre-active behavior in ONE place)"],
 ["EX_MSG_APPEND","Procedure call (exit class)","msg keys + vars → RES row","Uniform result append incl. WAITING_END"],
], widths=[1.5,1.3,1.7,2.0], font_size=8)

# ============ 5 DECISION TABLES ============
doc.add_page_break()
doc.add_heading("5. Decision Tables (content = business-maintainable)", level=1)
doc.add_heading("5.1 DT_RATECLASS_MATRIX", level=2)
table(["Rate Category (pattern)","Rate Class"], [
 ["RS*, RES*, E-1*, G-1*","RES"],
 ["GS*, COM*, C&I*, E-2*, LGS*","COM"],
 ["(fallback)","RES"],
], widths=[3.2,1.5])
doc.add_heading("5.2 DT_WAITING_PERIODS", level=2)
table(["Program","Prior Outcome","Waiting (months)","Message"], [
 ["CPP","D (de-enrolled)","6","045"],
 ["DPP","D","6","045"],
 ["PTR","D","3","045"],
 ["PPS","D","12","045"],
 ["AMP","F (default)","24","046"],
 ["DPA","F","12","046"],
 ["PREPAY","D","3","045"],
], widths=[1.1,1.5,1.4,1.0])
doc.add_heading("5.3 DT_PROGRAM_EXCLUSIONS (mutual exclusions)", level=2)
table(["Requested Program","Blocking Active Program","Severity","Message"], [
 ["PREPAY","MED (Medical Alert)","E","052"],
 ["PREPAY","BB (Budget Billing)","E","056"],
 ["PREPAY","DPA","E","057"],
 ["BB","PREPAY","E","056"],
 ["CSDD","PREPAY","E","055"],
 ["DPA","DPA (active plan)","E","012"],
 ["DPP","CPP","E","058 (one dynamic rate at a time)"],
 ["CPP","DPP","E","058"],
 ["TOU","DPP or CPP","W","340 (rate change replaces current dynamic rate)"],
 ["GREEN","CSOL","W","341 (stacking review)"],
], widths=[1.5,1.9,0.9,1.5])
doc.add_heading("5.4 Per-Program Rule Content (hard stops / warnings / not-yet-eligible)", level=2)
p("Complete rule content by program. Column 'P-mode' marks rules that EX_SEVERITY_MODE downgrades to "
  "I (not yet eligible) during Move-In pre-active evaluation. Message numbers refer to ZEC_ELIG (Section 7).")
table(["Program","Hard Stops (E)","Warnings (W)","P-mode"], [
 ["DPP Dynamic Peak Pricing","No AMI (044); RATE_CLASS≠RES (002); waiting period (045); active CPP (058)","—","—"],
 ["PTR Peak-Time Rebate","No AMI (044); AMI comm disabled (018); ≠RES (002); waiting (045)","—","—"],
 ["CPP Critical Peak Pricing","No AMI (044); ≠RES (002); waiting (045); active DPP (058)","—","—"],
 ["TOU Time-of-Use","No AMI (044)","Replaces active dynamic rate (340)","—"],
 ["PPS Peak Power Savers","≠RES (002); waiting (045)","—","Device install after activation (093)"],
 ["THERM Smart Thermostat","≠RES (002)","No Wi-Fi at premise (342)","—"],
 ["EVMC EV Managed Charging","—","Telematics compatibility unverified (320)","—"],
 ["EEREB EE Rebates","—","Invoice verification required (321)","—"],
 ["CSDD Due Date","Prepay active (055)","Returned payments ≥2 in 12M (031)","—"],
 ["EBILL Paperless","—","No valid e-mail on BP (322)","—"],
 ["DPA Payment Plan","Active plan (012); no qualifying arrears (072)","Broken plans ≥2 in 24M (330)","Contract/billing required (090)"],
 ["BB Budget Billing","Prepay active (056)","—","History <12 months (062)"],
 ["PREPAY Pay-As-You-Go","Medical flag (052); BB active (056); DPA active (057); ≠RES (002); no AMI (044)","—","—"],
 ["AUTOPAY","—","Returned payments ≥2 (031)","—"],
 ["PEXT Payment Extension","Extension already active (073)","≥2 extensions in 12M (331)","Contract required (090)"],
 ["SUMBILL Summary Billing","CA_COUNT_BP <2 (074)","—","—"],
 ["TPN Third-Party Notice","—","—","—"],
 ["LIHEAP","≠RES (002)","Income verification pending (332)","—"],
 ["PIPP","≠RES (002)","Income verification pending (332)","—"],
 ["AMP Arrearage Mgmt","No qualifying arrears (071); waiting after default (046)","—","Contract/billing required (090)"],
 ["CARE / FERA","≠RES (002)","Income self-certification post-check (332)","—"],
 ["MED Medical Alert","—","Physician certification due in 30 days (333)","—"],
 ["GREEN Green Power","—","Stacking with CSOL (341)","—"],
 ["CSOL Community Solar","—","Array capacity waitlist (334)","—"],
 ["ALERTS Usage/High-Bill","No AMI (044); AMI comm disabled (018)","—","—"],
 ["INTR Interruptible","≠COM (003); load < contract min (080)","—","—"],
 ["ADR OpenADR","≠COM (003)","Telemetry certification pending (335)","—"],
 ["CDR C&I Demand Response","≠COM (003); CURTAILABLE_LOAD_KW <50 (080)","—","—"],
], widths=[1.5,2.4,1.6,1.0], font_size=7.5)
doc.add_heading("5.5 DT_EXIT_RULES (de-enrollment)", level=2)
table(["Program","Exit Condition Checked","Severity","Message"], [
 ["DPA","Open installment balance > 0","E","EC_EXIT 007"],
 ["BB","Unbilled true-up pending at annual review date","W","EC_EXIT 010"],
 ["THERM / PPS","Open service order for device removal/install","E","EC_EXIT 011"],
 ["TOU","Within 12-month minimum participation","E","EC_EXIT 012"],
 ["INTR","Within contract term","E","EC_EXIT 013"],
 ["PREPAY","Negative balance outstanding","E","EC_EXIT 014"],
 ["All programs","(none matched)","—","exit allowed"],
], widths=[1.3,3.0,0.8,1.3])

# ============ 6 RULESETS ============
doc.add_page_break()
doc.add_heading("6. Rulesets", level=1)
p("28 rulesets assigned to F_EC_ELIGIBILITY_ALL, one per program, each with precondition "
  "PROGRAM_FILTER is empty OR contains <PROGRAM>. Standard internal sequence (worked example below applies "
  "the pattern to every program):")
numbered([
 "RU_<PGM>_010_EXCLUSIONS — EX_PROGRAM_EXCLUSION for the program vs T_ACTIVE; append result rows.",
 "RU_<PGM>_020_HARDSTOPS — evaluate DT_<PGM> hard-stop conditions per T_CA row (via loop); severity through EX_SEVERITY_MODE; append.",
 "RU_<PGM>_030_WAITING — EX_WAITING_PERIOD; append with WAITING_END.",
 "RU_<PGM>_040_WARNINGS — warning conditions; append severity W.",
 "RU_<PGM>_050_PREACTIVE — mode-P-only informational rows (093/090/062 family) where defined.",
])
doc.add_heading("6.1 Worked Example: RS_DPP_INT_VALIDATION (migrated legacy ruleset — Dynamic Peak Pricing)", level=2)
table(["Rule","Condition","Action"], [
 ["RU_DPP_010_EXCLUSIONS","T_ACTIVE contains CPP for same CA","Append E/058"],
 ["RU_DPP_020_AMI","AMI_METER = false (per CA)","Append E/044 (mode P: E stays E — device rules are premise-based and known pre-active)"],
 ["RU_DPP_021_RATECLASS","RATE_CLASS ≠ RES","Append E/002"],
 ["RU_DPP_030_WAITING","EX_WAITING_PERIOD(DPP) returns future date","Append E/045 with WAITING_END"],
 ["RU_DPP_900_LOG","always","Statistics update (exit class) — evaluation counter per program"],
], widths=[1.7,2.4,2.4], font_size=8.5)
p("Note: the legacy ruleset's device rule (legacy RU_DPP_CHECK_NOAIRDIV pattern) maps to RU_DPP_020/021; "
  "legacy rule numbering is preserved in the rule description field for traceability to the CE R5 "
  "specifications.")

# ============ 7 MESSAGES ============
doc.add_heading("7. Message Class ZEC_ELIG — Catalog", level=1)
table(["No.","Sev","Text (EN) — & = placeholder"], [
 ["002","E","Program &1 is restricted to residential rate class"],
 ["003","E","Program &1 is restricted to commercial rate class"],
 ["012","E","Active installment plan already exists on contract account &1"],
 ["018","E","AMI meter communication disabled — open meter ticket &1 must be resolved first"],
 ["044","E","Program &1 requires an AMI meter — premise has non-AMI device"],
 ["045","E","Re-enrollment waiting period active until &1 (previous participation ended &2)"],
 ["046","E","Waiting period after plan default active until &1"],
 ["052","E","Prepayment not permitted — active Medical Alert / life-support equipment at premise"],
 ["055","E","Not compatible with active Prepay service"],
 ["056","E","Prepay and Budget Billing are mutually exclusive"],
 ["057","E","Prepay requires settlement of the active payment plan first"],
 ["058","E","Only one dynamic rate program allowed — &1 is active"],
 ["062","I","Not yet eligible — requires 12 months billing history at premise (&1 of 12)"],
 ["071","E","No eligible arrears — program requires a qualifying past-due balance"],
 ["072","E","No qualifying arrears for a payment plan"],
 ["073","E","A payment extension is already active until &1"],
 ["074","E","Summary billing requires at least two contract accounts"],
 ["080","E","Curtailable load &1 kW below program minimum &2 kW"],
 ["090","I","Not yet eligible — requires an active contract with billing history"],
 ["093","I","Device installation is scheduled after contract activation"],
 ["031","W","&1 returned payments in last 12 months — supervisor approval required"],
 ["320","W","Vehicle/charger telematics compatibility not yet verified"],
 ["321","W","Rebate subject to invoice verification"],
 ["322","W","No valid e-mail address on business partner"],
 ["330","W","&1 defaulted payment plans in last 24 months"],
 ["331","W","&1 payment extensions already granted in last 12 months"],
 ["332","W","Income verification pending — enrollment provisional"],
 ["333","W","Physician certification required within 30 days"],
 ["334","W","Community solar array at capacity — waitlist position assigned"],
 ["335","W","Interval telemetry certification pending"],
 ["340","W","Enrollment replaces the active dynamic rate program &1"],
 ["341","W","Review stacking with active subscription &1"],
 ["342","W","No Wi-Fi at premise — installation survey required"],
], widths=[0.6,0.5,5.4], font_size=8)

# ============ 8 INVOCATION ============
doc.add_page_break()
doc.add_heading("8. Invocation from Z_EC_ELIGIBILITY_PRECHECK (ABAP skeleton)", level=1)
for line in [
 "DATA(lo_fct) = cl_fdt_factory=>if_fdt_factory~get_instance( )->get_function( iv_id = gc_f_elig_all_id ).",
 "DATA(lo_ctx) = lo_fct->get_process_context( ).",
 "lo_ctx->set_value( iv_name = 'BUSINESS_PARTNER' ia_value = iv_partner ).",
 "lo_ctx->set_value( iv_name = 'MODE'             ia_value = iv_mode ).      \" S / P",
 "lo_ctx->set_value( iv_name = 'T_CA'             ia_value = lt_ctx_ca ).    \" pre-fetched, parallel aRFC",
 "lo_ctx->set_value( iv_name = 'T_HISTORY'        ia_value = lt_hist ).",
 "lo_ctx->set_value( iv_name = 'T_ACTIVE'         ia_value = lt_active ).",
 "lo_fct->process( EXPORTING io_context = lo_ctx IMPORTING eo_result = DATA(lo_res) ).",
 "lo_res->get_value( IMPORTING ea_value = et_result ).   \" RES_EC_ELIG",
 "\" et_result returned to BTP via RFC; SLG1 log written with context hash + row count.",
]:
    mono(line)
bullets([
 "Function IDs read once from ZEC_FDT_IDS constants table (filled at transport import) — never hardcoded GUIDs.",
 "Generation: ensure 'Generated Code' is active; first call after import triggers generation — the post-import job (Config Guide 9.4) pre-generates to avoid first-user latency.",
])

# ============ 9 TESTING ============
doc.add_heading("9. Simulation and Testing", level=1)
bullets([
 ("Workbench simulation","every ruleset ships with saved simulation variants (BRFplus workbench → Function → Simulation): one per program covering eligible, each hard stop, each warning, and mode P."),
 ("Regression suite","eCATT/ABAP Unit harness ZEC_TEST_ELIGIBILITY replays the legacy DPP enrollment test-case library (Nov/Dec test workbooks in the project repository) against F_EC_ELIGIBILITY_ALL and compares result tables."),
 ("Trace","FDT trace (technical trace on function) activated in QA for defect analysis; lean trace in PRD sampled 1%."),
 ("Acceptance gate","rule content changes require: simulation variants green + regression suite green + business owner approval — enforced by the change process, not the tool."),
])

# ============ 10 TRANSPORT ============
doc.add_heading("10. Transport, Versioning and Maintenance Governance", level=1)
bullets([
 ("Transport","structural objects (application, functions, data objects, expressions, rulesets) via workbench transports; decision-table CONTENT transported by default with the object — for business-maintained tables (Section 5) switch content to 'local/productive maintenance' only if the client accepts content divergence; recommended: content changes in DEV + transport, emergency content changes in PRD with retro-transport."),
 ("Versioning","BRFplus object versioning ON for decision tables and rulesets; version comment mandatory (change request reference)."),
 ("Post-import","job FDT_TRANS / regeneration (Config Guide 9.4) scheduled after each import."),
 ("Maintenance roles","business rule owners: change access limited to DT_* content via authorization (S_FDT with object-level restriction + application exit authorization hook); developers: full application in DEV only."),
 ("Auditability","BRFplus change log + version history reviewed in the monthly governance meeting (CS-02 pattern); eligibility statistics (rule 900 counters) feed the program adoption dashboard."),
])


# ============ 11 OBJECT INVENTORY ============
doc.add_heading("11. Object Inventory and Consolidated Implementation Model", level=1)
p("Definitive object counts for build planning. The per-program condition content of Section 5.4 is "
  "implemented in a CONSOLIDATED model: criteria-matrix decision tables carry one ROW per program instead of "
  "one table per program. The naive per-program alternative (DT_<PGM>_HARDSTOPS / DT_<PGM>_WARNINGS) would "
  "require 40 decision tables; the consolidated model requires 7 and concentrates business maintenance in two "
  "matrices. Productive content for every table ships in the companion workbook "
  "'Enrollment Center – BRFplus Configuration Workbook v1.0.xlsx' (one sheet per decision table, structured "
  "for BRFplus Workbench Excel import).")
table(["Object Type","Count","Objects"], [
 ["Application","1","ZEC_PROGRAM_ELIGIBILITY"],
 ["Functions","2","F_EC_ELIGIBILITY_ALL (event mode); F_EC_DEENROLL_CHECK (functional mode)"],
 ["Rulesets","28","RS_<PGM>_VALIDATION per program (RS_DPP_INT_VALIDATION retains its legacy name)"],
 ["Decision tables","7","DT_RATECLASS_MATRIX; DT_PROGRAM_CRITERIA (criteria matrix, 28 rows, drives all generic hard stops); DT_PROGRAM_EXCLUSIONS; DT_WAITING_PERIODS; DT_PROGRAM_WARNINGS (16 warning rows); DT_PREACTIVE_INFO; DT_EXIT_RULES"],
 ["Reusable expressions","7","EX_RATE_CLASS, EX_WAITING_PERIOD, EX_PROGRAM_EXCLUSION, EX_AMI_CHECK, EX_HIST_MONTHS_OK, EX_SEVERITY_MODE, EX_MSG_APPEND"],
 ["Data objects","4","CTX_EC_ELIG, CTX_EC_CA, CTX_EC_HIST, RES_EC_ELIG"],
 ["Message class","1 (33 messages)","ZEC_ELIG"],
], widths=[1.6,1.0,3.9], font_size=8.5)
p("Mapping note: rule RU_<PGM>_020_HARDSTOPS evaluates the program's row of DT_PROGRAM_CRITERIA through a "
  "shared generic expression — each violated criterion appends its fixed ZEC_ELIG message (rate class → "
  "002/003, AMI → 044, AMI comm → 018, prepay/medical/BB/DPA blocks → 052–057, arrears → 071/072, history → "
  "062, load → 080, CA count → 074, extension → 073). Program-specific one-off conditions that do not fit the "
  "matrix (e.g., dynamic-rate mutual exclusion 058) remain in DT_PROGRAM_EXCLUSIONS.")

out = r"C:\Users\jnamm\OneDrive\Desktop\Service Cloud for Utilities\Enrollment Center\Enrollment Center - BRFplus Detailed Configuration v1.0.docx"
doc.save(out)
print("Saved:", out)
