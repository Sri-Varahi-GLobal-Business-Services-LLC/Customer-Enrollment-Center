# -*- coding: utf-8 -*-
"""Generate Enrollment Center Runtime Service Interactions and Data Exchange document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT="1F4E79"; HEADER_FILL="1F4E79"; ALT_FILL="DEEAF6"
doc = Document()
n = doc.styles["Normal"]; n.font.name="Arial"; n.font.size=Pt(10.5); n.paragraph_format.space_after=Pt(6)
for lvl, sz in [("Heading 1",16),("Heading 2",13),("Heading 3",11.5)]:
    st = doc.styles[lvl]; st.font.name="Arial"; st.font.size=Pt(sz); st.font.bold=True
    st.font.color.rgb = RGBColor(0x1F,0x4E,0x79)
    rpr = st.element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
    if rf is None: rf = OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a),'Arial')

def p(text, bold=False, size=None, color=None, align=None):
    para = doc.add_paragraph(); r = para.add_run(text); r.bold=bold
    if size: r.font.size=Pt(size)
    if color: r.font.color.rgb=RGBColor.from_string(color)
    if align: para.alignment=align
    return para

def mono(text):
    para = doc.add_paragraph(); r = para.add_run(text)
    r.font.name="Consolas"; r.font.size=Pt(8.5)
    para.paragraph_format.space_after=Pt(1)
    return para

def bullets(items):
    for it in items:
        if isinstance(it, tuple):
            para = doc.add_paragraph(style="List Bullet")
            r = para.add_run(it[0]+" — "); r.bold=True
            para.add_run(it[1])
        else:
            doc.add_paragraph(it, style="List Bullet")

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def table(headers, rows, widths=None, font_size=8):
    t = doc.add_table(rows=1, cols=len(headers)); t.style="Table Grid"
    t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""
        r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(font_size)
        r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); shade(c,HEADER_FILL)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=""
            run=cells[i].paragraphs[0].add_run(str(val)); run.font.size=Pt(font_size)
            if ri%2==1: shade(cells[i],ALT_FILL)
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

def action(hid, title, trigger, steps, notes=None):
    doc.add_heading(hid + " " + title, level=2)
    p("Trigger: " + trigger, bold=False)
    table(["#","From → To","Service / Interface","Sync","Data Exchanged (request ⇢ / response ⇠)"], steps,
          widths=[0.35,1.25,1.6,0.5,3.15])
    if notes: bullets(notes)

# ============ COVER ============
for _ in range(6): doc.add_paragraph()
p("Enrollment Center — Runtime Service Interactions", bold=True, size=25, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER)
p("Services Required and Data Exchange per User Action", bold=True, size=17, color="404040", align=WD_ALIGN_PARAGRAPH.CENTER)
p("SAP Service Cloud V2 ↔ SAP BTP ↔ SAP S/4HANA Utilities — action-by-action sequence specification",
  size=12, color="595959", align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(7): doc.add_paragraph()
table(["Attribute","Value"], [
 ["Document","Runtime Service Interactions & Data Exchange"],
 ["Version / Status","1.1 (per-utility device model, DER assets in context, combo CA patterns)"],
 ["Date","08 July 2026"],
 ["Related Documents","Solution Design v2.0 (flows, status model); TDD v1.1 (APIs, CDS, tables); Config Guide v1.0 (destinations, iFlows); BRFplus Config v1.1"],
 ["Audience","Integration developers, CAP developers, S/4HANA developers, testers"],
], widths=[1.7,4.8], font_size=9)
doc.add_page_break()

# ============ 1 SERVICE REGISTRY ============
doc.add_heading("1. Service Registry (all services used at runtime)", level=1)
doc.add_heading("1.1 SAP Service Cloud V2 services", level=2)
table(["ID","Service","Direction","Purpose"], [
 ["V2-01","External application / mashup runtime (signed context token)","SC V2 → UI","Launch EC with agent/customer/case (or move-in) context; token carries agentId, businessPartnerId, contractAccountId?, caseId | moveInId, mode, exp"],
 ["V2-02","Case Service REST (POST /case-service/cases/{id}/timeline)","BTP → SC V2","Write enrollment status entries to the case timeline"],
 ["V2-03","Case Service REST (PATCH case)","BTP → SC V2","Optional: set case fields (follow-up flag on rejection)"],
 ["V2-04","Identity (IAS OIDC)","UI ↔ IAS","Silent SSO of the agent session into the embedded UI"],
], widths=[0.6,2.4,1.0,2.9])
doc.add_heading("1.2 BTP Enrollment Service APIs (consumed by the UI)", level=2)
table(["ID","API","Purpose"], [
 ["EC-01","GET /enrollment/v1/context/{bp}?mode=&moveInId=","Aggregated customer context + participation + history"],
 ["EC-02","POST /enrollment/v1/eligibility:evaluate","All-programs eligibility (single round trip)"],
 ["EC-03","GET /enrollment/v1/programs[?programId=]","Catalog + capture form metadata + terms"],
 ["EC-04","POST /enrollment/v1/enrollments (Idempotency-Key)","Create enrollment request"],
 ["EC-05","GET /enrollment/v1/enrollments/{id} | ?bp= | ?moveInId=","Status read / lists"],
 ["EC-06","POST /enrollment/v1/enrollments/{id}:requestOverride","Warning override to supervisor"],
 ["EC-07","POST /enrollment/v1/deenrollments:validate | POST /deenrollments","Exit check / de-enrollment"],
 ["EC-08","POST /enrollment/v1/enrollments/{id}/documents","Consent/supporting document"],
 ["EC-09","GET /enrollment/v1/exceptions?status=open","Supervisor exception work list"],
], widths=[0.7,3.3,2.9])
doc.add_heading("1.3 S/4HANA services (via Integration Suite / destinations)", level=2)
table(["ID","Service","Protocol","Consumed by"], [
 ["S4-01","Z_EC_CONTEXT_SRV (consolidated OData over I_Utilities* CDS + FI-CA facts)","OData V4, principal propagation","IF_EC_Context_Read"],
 ["S4-02","API_BUSINESS_PARTNER","OData V2","IF_EC_Context_Read"],
 ["S4-03","Z_EC_ELIGIBILITY_PRECHECK (BRFplus wrapper)","RFC","IF_EC_Eligibility"],
 ["S4-04","Z_EC_* fulfillment wrappers (RATECHANGE, INSTPLAN, BUDGETBILL, CAUPDATE, PAYMENT, ASSIST, CORR, BILLGROUP, SERVICEORDER)","RFC","IF_EC_Fulfillment_Dispatch"],
 ["S4-05","Z_EC_DEENROLL_CHECK + reversal wrappers","RFC","IF_EC_Deenroll"],
 ["S4-06","Z_EC_ENROLLMENT_SRV (ZEC_ENROLL_HDR/_HIST)","OData","CAP (participation write-through) / reconciliation"],
 ["S4-07","Event enablement → Event Mesh topics utilco/s4/…","Events (AMQP)","IF_EC_Status_Events, IF_EC_MoveIn_Events"],
 ["S4-08","ZEC_RECON_EXTRACT","OData","Reconciliation job"],
], widths=[0.7,3.1,1.4,1.7])
p("Direction of truth per exchange: reads federate live from S/4HANA (no staging); writes always land in "
  "S/4HANA as business objects; BTP persists process state (EnrollmentRequest) and participation is "
  "written through to ZEC_ENROLL_HDR so S/4HANA remains system of record for participation.")

# ============ 2 ACTIONS ============
doc.add_page_break()
doc.add_heading("2. Data Exchange per User Action", level=1)

action("A1.","Agent opens Enrollment Center (work center / Customer 360)",
 "Agent clicks the Enrollment Center work center for an identified, verified customer.",
 [
  ["1","SC V2 → UI","V2-01 mashup launch","sync","⇢ signed context: agentId, businessPartnerId, contractAccountId (opt), caseId, mode=STANDALONE, exp 300 s"],
  ["2","UI → IAS","V2-04 OIDC","sync","Silent auth; ⇠ JWT with ec.Agent scope"],
  ["3","UI → CAP","EC-01 GET /context/{bp}","sync","⇢ bp, mode ⇠ header (name, verified), T_CA[] (CA, division, premise, contract, rate, AMI flags, financial flags), participation[], history[]"],
  ["4","CAP → IF_EC_Context_Read → S/4","S4-01 + S4-02 (parallel $batch)","sync","⇢ bp ⇠ BP header; per CA x division: FKKVKP facts, EVER contract, EANLH rate, premise, installed device (device no., meter no., AMI/AMR type, connection status) incl. I_UtilsAdvncdMeteringSystem, dunning level, returned payments, arrears, locks; DER assets fetched in parallel from the DER platform API"],
  ["5","CAP → HANA Cloud","internal","sync","⇠ in-flight EnrollmentRequests (status not final) merged over S4-06 participation"],
  ["6","UI → CAP","EC-02 POST /eligibility:evaluate","sync","⇢ bp, T_CA keys, mode=S ⇠ RES_EC_ELIG rows: per program/CA eligible|E|W|I + message text + waiting end"],
  ["7","CAP → IF_EC_Eligibility → S/4","S4-03 RFC","sync","⇢ CTX_EC_ELIG (pre-fetched context) ⇠ RES_EC_ELIG; CAP caches for the interaction (TTL 5 min, invalidated on any write)"],
  ["8","UI","render","—","Four panels render complete on first paint (<3 s target)"],
 ],
 [("Failure behavior","step 4/7 timeout → UI renders participation from CAP store with banner 'live data unavailable', enroll actions disabled; retry button re-invokes EC-01/EC-02.")])

action("A2.","Agent opens Enrollment Center inside Move-In (embedded step)",
 "Move-In guided flow reaches the Customer Programs step (contract not yet active).",
 [
  ["1","SC V2 → UI","V2-01 launch (move-in variant)","sync","⇢ context: agentId, businessPartnerId (new or existing), moveInId, prospective premiseId/installation, requestedStart, mode=MOVE_IN"],
  ["2","UI → CAP","EC-01 GET /context/{bp}?mode=P&moveInId=","sync","⇠ prospective context: premise/installation/device from move-in data; participation empty or existing-BP participation; history for the BP"],
  ["3","CAP → IF_EC_Context_Read → S/4","S4-01 (premise/installation/device by moveInId keys)","sync","⇢ premise/installation keys ⇠ device/AMI capability of the target premise, premise type"],
  ["4","UI → CAP → S/4","EC-02 → S4-03 with MODE=P","sync","⇠ eligibility where PREACTIVE=N programs omitted, I-rows from DT_PREACTIVE_INFO, history-based E downgraded to I (EX_SEVERITY_MODE)"],
 ],
 [("Key difference","CONTRACT_ACTIVE=false in every CTX_EC_CA row; nothing is written to S/4HANA in this action.")])

action("A3.","Agent opens program information (hyperlink)",
 "Agent clicks a program name in Available or Participating panel.",
 [
  ["1","UI → CAP","EC-03 GET /programs?programId=","sync","⇠ catalog row: name, category, terms text, recipeType, captureForm JSON (field defs incl. types/options), consent statement ID + version"],
 ],
 [("No backend call","catalog is CAP-owned (HANA Cloud), cached in the UI; S/4HANA is not touched.")])

action("A4.","Agent submits an enrollment (standalone, no warnings)",
 "Agent completes the capture form and clicks Submit Enrollment.",
 [
  ["1","UI → CAP","EC-04 POST /enrollments","sync","⇢ Idempotency-Key, bp, ca, programId, capture payload, consentStatementId, caseId ⇠ 201: requestId, status=Enrollment Request"],
  ["2","CAP → IF_EC_Eligibility → S/4","S4-03 (re-validation, single program)","sync","Defense against stale UI: hard stop now → 422 with RES rows, nothing persisted beyond audit"],
  ["3","CAP → HANA Cloud","internal","sync","EnrollmentRequest persisted (status Enrollment Request); AuditLog append (actor, payload hash)"],
  ["4","CAP → Event Mesh","publish utilco/ec/fulfillment/dispatch","async","⇢ requestId, recipeType, target keys, capture payload"],
  ["5","CAP → IF_EC_SCV2_Timeline → SC V2","V2-02","async","⇢ caseId, text 'Enrollment <program> submitted', requestId reference"],
  ["6","Event Mesh → IF_EC_Fulfillment_Dispatch → S/4","S4-04 (wrapper per recipeType)","async","⇢ e.g. RECIPE_RATE: installation, target rate, effective date | RECIPE_DPA: CA, open items selection, plan params ⇠ backend keys (change doc / plan no. / order no.) or error"],
  ["7","S/4 wrapper → ZEC_ENROLL_HDR","internal (S4-06 write)","sync in wrapper","Participation row created/updated (status pending) + change doc"],
  ["8","S/4 → Event Mesh","S4-07 utilco/s4/enrollment/confirmed","async","⇢ requestId, backend keys, actualDate, status ENROLLED | REJECTED(reason)"],
  ["9","Event Mesh → IF_EC_Status_Events → CAP","EventInbox","async","Dedup by eventId; EnrollmentRequest → Enrolled (actualDate) | Enrollment Rejected"],
  ["10","CAP → SC V2","V2-02 timeline","async","'<program> Enrolled effective <date>' — agent sees it without reopening the EC"],
  ["11","UI (open) → CAP","EC-05 poll/refresh on panel focus","sync","⇠ current status chip transitions Request → Pending → Enrolled"],
 ])

action("A5.","Enrollment with warning-level failure (supervisor override)",
 "Program shows WARNING; agent clicks Request supervisor approval.",
 [
  ["1","UI → CAP","EC-06 POST :requestOverride","sync","⇢ programId, messageCode, justification ⇠ 202; request persisted status Pending Approval"],
  ["2","CAP → SBPA","process API EC_SupervisorOverride","async","⇢ context (customer, program, warning text, justification, agent)"],
  ["3","SBPA → Supervisor inbox","workflow UI","—","Supervisor approves / rejects with reason"],
  ["4","SBPA → CAP","callback API","sync","⇠ decision; OverrideRecord persisted (approver, reason, timestamp)"],
  ["5","CAP","continue as A4 steps 2–11 (approve) or status Enrollment Rejected + timeline (reject)","mixed","Audit links OverrideRecord to the request"],
 ])

action("A6.","Enrollment submitted inside Move-In (pre-active)",
 "Agent submits a program in the embedded Customer Programs step.",
 [
  ["1","UI → CAP","EC-04 with moveInId, mode=P","sync","⇠ 201 status Enrollment Request; request PARKED (no dispatch event published)"],
  ["2","CAP → HANA Cloud","internal","sync","EnrollmentRequest stored with moveInId binding; audit"],
  ["3","CAP → SC V2","V2-02 (move-in case timeline)","async","'Program <x> captured — fulfillment on contract activation'"],
 ],
 [("Nothing reaches S/4HANA yet","the only S/4HANA writes happen after the activation event (A7).")])

action("A7.","Move-In completed in S/4HANA (deferred fulfillment release)",
 "Back office / batch completes the move-in; contract becomes active.",
 [
  ["1","S/4 → Event Mesh","S4-07 utilco/s4/movein/completed","async","⇢ moveInId, bp, new CA/contract keys, activation date"],
  ["2","Event Mesh → IF_EC_MoveIn_Events → CAP","EventInbox","async","Parked requests selected via EC-05 ?moveInId="],
  ["3","CAP → S/4","S4-03 re-validation per request (mode S, now-active CA)","sync","Hard stop now → Enrollment Rejected + timeline + follow-up task (V2-03)"],
  ["4","CAP → Event Mesh","utilco/ec/fulfillment/dispatch per surviving request","async","Continues exactly as A4 steps 6–11 with the new contract account keys"],
 ])

action("A8.","Move-In cancelled in S/4HANA (cascade)",
 "Move-in document is cancelled before activation.",
 [
  ["1","S/4 → Event Mesh → CAP","utilco/s4/movein/cancelled via IF_EC_MoveIn_Events","async","⇢ moveInId"],
  ["2","CAP","internal","sync","All parked requests (moveInId) → status Cancelled; audit rows"],
  ["3","CAP → SC V2","V2-02","async","Timeline: 'n enrollment request(s) cancelled with move-in'"],
 ])

action("A9.","Agent de-enrolls a customer from a program",
 "Agent clicks De-enroll on a Participating Programs row.",
 [
  ["1","UI → CAP","EC-07 POST /deenrollments:validate","sync","⇢ enrollmentId ⇠ exit-rule result (S4-05 Z_EC_DEENROLL_CHECK via IF_EC_Deenroll): blocked (E + EC_EXIT msg) or clear (+ W info)"],
  ["2","UI","reason capture","—","Agent selects reason code (blocked case ends here with message display)"],
  ["3","UI → CAP","EC-07 POST /deenrollments","sync","⇢ enrollmentId, reasonCode ⇠ status Pending De-enrollment"],
  ["4","CAP → Event Mesh → IF → S/4","dispatch reversal recipe (S4-04/S4-05)","async","e.g. rate change back to standard / plan deactivation / order for device removal; ZEC_ENROLL_HDR end-dated, HIST written"],
  ["5","S/4 → Event Mesh → CAP → SC V2","utilco/s4/enrollment/confirmed + V2-02","async","Status De-enrolled; end date in Program History; waiting-period rules now apply automatically at next eligibility call"],
 ])

action("A10.","Fulfillment failure and exception handling",
 "A dispatched recipe fails in S/4HANA (business error or technical fault).",
 [
  ["1","IF_EC_Fulfillment_Dispatch","retry policy","async","3 retries exponential (technical faults); business errors (wrapper E messages) skip retry"],
  ["2","IF → DLQ → IF_EC_Exception","EC_DLQ","async","⇢ requestId, error payload; EnrollmentRequest → Error – In Review"],
  ["3","IF_EC_Exception → SBPA","EC_ExceptionReview process","async","Supervisor task with context + actions"],
  ["4","(where BPEM used)","S/4 wrapper opened BPEM case ZEC1","—","Case number linked on the request; utilco/s4/bpem/casechanged closes the loop on resolution"],
  ["5","Supervisor action → CAP","SBPA callback: retry | correct+resubmit | cancel","sync","Retry re-publishes dispatch; cancel → status + timeline + callback task on case (V2-03)"],
 ])

action("A11.","Reconciliation (scheduled, no user action)",
 "Daily CAP job — safety net for missed events.",
 [
  ["1","CAP → S/4","S4-08 ZEC_RECON_EXTRACT","sync","⇠ participation + in-flight fulfillment states since last run"],
  ["2","CAP","compare vs EnrollmentRequests","—","Pending Backend older than SLA without backend state → exception (A10 path); mismatched statuses corrected + audited"],
 ])

# ============ 3 PAYLOADS ============
doc.add_page_break()
doc.add_heading("3. Key Payload Structures (abbreviated)", level=1)
p("EC-01 context response:", bold=True)
mono('{ "bp":"1000004711", "name":"…", "mode":"S", "cas":[ { "ca":"300000123456", "division":"01",')
mono('   "premise":{...}, "contract":{"id":"…","active":true,"rateCategory":"RS_1000"},')
mono('   "device":{"deviceNo":"10004481","meterNo":"1KA-88213045","type":"AMI","status":"CONNECTED",')
mono('             "model":"…","lastReadDate":"2026-07-05","lastReadValue":41208},')
mono('   "fica":{"dunningLevel":1,"arrears":342.50,')
mono('   "returnedPayments12m":0,"activeInstallmentPlan":true,"budgetBilling":true,"prepay":false,')
mono('   "medicalFlag":false,"billingHistoryMonths":26} } ],')
mono('  "derAssets":[{"type":"EV_CHARGER_L2","capacityKw":11,"status":"ACTIVE"}],')
mono('  "participation":[...], "history":[...] }')
p("Note: cas[] carries one entry per contract account x division — both combo patterns (CA per utility, or a single CA spanning divisions) are represented by this shape; device data is per division. derAssets come from the DER platform (empty array when unavailable).", bold=False)
p("EC-02 eligibility response row:", bold=True)
mono('{ "programId":"CSDD", "ca":"300000123456", "eligible":"N", "severity":"W",')
mono('  "msg":{"id":"ZEC_ELIG","no":"031","text":"2 returned payments…"}, "waitingEnd":null }')
p("utilco/ec/fulfillment/dispatch:", bold=True)
mono('{ "eventId":"…", "requestId":"REQ-1001", "recipeType":"RECIPE_RATE",')
mono('  "target":{"installation":"4000123","ca":"300000123456"},')
mono('  "payload":{"targetRate":"RS_1007","effectiveDate":"2026-08-01"}, "agentId":"ARIVERA" }')
p("utilco/s4/enrollment/confirmed:", bold=True)
mono('{ "eventId":"…", "requestId":"REQ-1001", "status":"ENROLLED", "actualDate":"2026-08-01",')
mono('  "backendKeys":{"changeDoc":"…","planNo":null,"orderNo":null,"bpemCase":null} }')

# ============ 4 SLA ============
doc.add_heading("4. Interaction SLAs and Modes Summary", level=1)
table(["Exchange","Mode","Target"], [
 ["EC-01 + EC-02 (open EC, full render)","sync","< 3 s end-to-end"],
 ["EC-04 submit acknowledgment","sync","< 1.5 s (re-validation included)"],
 ["Fulfillment dispatch → Enrolled (simple recipes: CSDD, EBILL, AUTOPAY)","async","< 1 min typical"],
 ["Rate-change recipes (DPP/CPP/TOU…)","async","minutes; effective next bill cycle (business date)"],
 ["Timeline write-back after status change","async","< 30 s"],
 ["Move-in activation → parked request dispatch","async","< 2 min from event"],
 ["Reconciliation detection of missed event","batch","≤ 24 h"],
], widths=[3.4,0.9,2.2], font_size=9)
p("Design invariants: every write path is idempotent (Idempotency-Key / eventId dedup); the agent is never "
  "blocked by an async leg — sync legs end at the acknowledgment, and all later state reaches the agent via "
  "the SC V2 case timeline and the status chips on re-render.")

out = r"C:\Users\jnamm\OneDrive\Desktop\Service Cloud for Utilities\Enrollment Center\Enrollment Center - Runtime Service Interactions and Data Exchange v1.0.docx"
doc.save(out)
print("Saved:", out)
