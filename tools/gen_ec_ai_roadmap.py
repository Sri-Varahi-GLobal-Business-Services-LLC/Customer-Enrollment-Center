# -*- coding: utf-8 -*-
"""Generate Enrollment Center AI/ML Roadmap document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT="1F4E79"; HEADER_FILL="1F4E79"; ALT_FILL="DEEAF6"
doc = Document()
n = doc.styles["Normal"]; n.font.name="Arial"; n.font.size=Pt(10.5); n.paragraph_format.space_after=Pt(6)
for lvl, sz in [("Heading 1",16),("Heading 2",13),("Heading 3",11.5)]:
    st = doc.styles[lvl]; st.font.name="Arial"; st.font.size=Pt(sz); st.font.bold=True
    st.font.color.rgb = RGBColor(0x1F,0x4E,0x79)
    rpr = st.element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
    if rf is None: rf = OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a),'Arial')

def p(text, bold=False, size=None, color=None, align=None):
    para = doc.add_paragraph(); r = para.add_run(text); r.bold=bold
    if size: r.font.size=Pt(size)
    if color: r.font.color.rgb=RGBColor.from_string(color)
    if align: para.alignment=align
    return para

def bullets(items):
    for it in items:
        if isinstance(it, tuple):
            para = doc.add_paragraph(style="List Bullet")
            r = para.add_run(it[0]+" — "); r.bold=True
            para.add_run(it[1])
        else:
            doc.add_paragraph(it, style="List Bullet")

def numbered(items):
    for it in items: doc.add_paragraph(it, style="List Number")

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def table(headers, rows, widths=None, font_size=8.5):
    t = doc.add_table(rows=1, cols=len(headers)); t.style="Table Grid"
    t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""
        r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(font_size)
        r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); shade(c,HEADER_FILL)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=""
            run=cells[i].paragraphs[0].add_run(str(val)); run.font.size=Pt(font_size)
            if ri%2==1: shade(cells[i],ALT_FILL)
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

# ===== COVER =====
for _ in range(6): doc.add_paragraph()
p("Customer Enrollment Center", bold=True, size=26, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER)
p("AI / ML Roadmap", bold=True, size=18, color="404040", align=WD_ALIGN_PARAGRAPH.CENTER)
p("Machine learning and generative AI on top of the Enrollment Center — use cases, SAP technology mapping, governance and phasing",
  size=12, color="595959", align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(7): doc.add_paragraph()
table(["Attribute","Value"], [
 ["Document","AI/ML Roadmap (extends Solution Design v2.0, Appendix B)"],
 ["Version / Status","1.1 (adds Section 3: propensity score calculation)"],
 ["Date","13 August 2026"],
 ["Governing principle","ML recommends — BRFplus decides. Hard eligibility remains deterministic and auditable."],
 ["Audience","Business stakeholders, architects, data science / AI teams"],
], widths=[1.7,4.8], font_size=9)
doc.add_page_break()

# ===== 1 =====
doc.add_heading("1. Why the Enrollment Center Is AI-Ready", level=1)
p("The Enrollment Center architecture creates, as a by-product of normal operation, exactly the assets AI/ML "
  "needs — none of them require new construction:")
bullets([
 ("Labeled outcome data","every enrollment attempt, override, rejection, de-enrollment and default lands in EnrollmentRequest / ZEC_ENROLL_HIST / OverrideRecord / AuditLog with full feature context — a growing training set from day one."),
 ("A single decision point","all programs flow through one eligibility call (EC-02) and one submission API (EC-04) — one place to attach scoring, ranking and explanation."),
 ("Structured knowledge","the program catalog (terms, criteria, capture forms) is machine-readable JSON — direct grounding material for generative AI."),
 ("Rich customer context","per-utility device (AMI/AMR, status), FI-CA financial facts, DER assets and interval data via the S/4HANA CDS layer."),
 ("A rules backbone","BRFplus provides the deterministic decision layer AI must never replace — which makes adding AI on top governable instead of risky."),
])

# ===== 2 USE CASES =====
doc.add_heading("2. Use Case Portfolio", level=1)
table(["#","Use Case","What It Does","Data","SAP Technology","Phase"], [
 ["AI-1","Next-best-program recommendation","Ranks Available Programs by enrollment propensity x value; 'Recommended' badge with reason; agent pitches the right program first","Program history, usage, device/DER, FI-CA facts, demographics","HANA Cloud PAL/APL (in-database scoring) or SAP AI Core; served via EC-10 scoring API enriching EC-02","1"],
 ["AI-2","Grounded agent copilot","Explains any program for THIS customer in plain language; objection handling; converts ZEC_ELIG codes into conversational explanations with remediation ('AMR meter - offer AMI exchange, then re-check')","Program catalog JSON, tariff documents, eligibility results","Generative AI Hub on SAP AI Core + Joule Studio skill; SC V2 CX AI Toolkit for call summarization","1"],
 ["AI-3","Bill-impact simulation","'Will TOU/DPP save this customer money?' - recompute 12 months of interval data under the target rate before enrolling; recommend only when beneficial","Interval/EDM data, rate definitions","Deterministic recompute on HANA Cloud + ML load-profile estimation for customers without full interval history","2"],
 ["AI-4","Payment-plan default risk","Predicts DPA/AMP plans likely to break; drives proactive outreach before dunning","FI-CA payment behavior, broken-plan history, arrears trajectory","HANA Cloud PAL classification; scores surfaced to collections and the supervisor cockpit","2"],
 ["AI-5","Assistance document intelligence","Auto-extracts income proof for LIHEAP/PIPP/CARE from uploaded documents; clears msg-332 verification backlog","EC-08 document uploads","SAP Document Information Extraction (BTP service)","2"],
 ["AI-6","Exception triage learning","Classifies Error-In-Review causes and suggests resolution; learns from supervisor actions","EventInbox, AuditLog, OverrideRecord (self-labeling)","SAP AI Core classification; SBPA task enrichment","3"],
 ["AI-7","DER/DR performance forecasting","Customer baseline (CBL) estimation for PTR settlement; thermostat/battery fleet yield forecasts for dispatch","Interval data, event history, weather","DER platform side (architecture group 6) - AI Core time-series; outside EC scope, consumes EC participation data","3"],
], widths=[0.4,1.35,2.5,1.5,2.2,0.45], font_size=7.5)

# ===== 3 GOVERNANCE =====
doc.add_heading("3. How the Recommendation Score (AI-1) Is Calculated", level=1)
p("The percentage shown next to a recommended program is a CALIBRATED ENROLLMENT PROBABILITY: "
  "P(customer enrolls in program X if offered). In the interactive demo the scores are illustrative seed "
  "values previewing the experience; in production they are produced by the following pipeline.")
table(["Step","What Happens","Detail"], [
 ["1. Training data","The Enrollment Center labels its own examples","Every display of the Available panel plus the outcome (enrolled within 30 days: yes/no) forms one example — accumulated in EnrollmentRequest / ZEC_ENROLL_HIST plus AI_RECO_shown/accepted telemetry from day one"],
 ["2. Features","Same context already fetched for EC-01","Device type and status, DER assets, arrears and payment behavior, billing history length, current participation, program history (prior de-enrollments), usage aggregates from interval data, rate class, tenure, season"],
 ["3. Model","Classification","Gradient boosting or logistic regression via HANA Cloud APL/PAL (in-database, automated); per-program models or one multi-program model with the program as a feature; output = raw probability"],
 ["4. Calibration","Raw probability becomes a trustworthy %","Platt scaling / isotonic regression so that ~87% of customers scored 87% actually enroll when offered — the property that makes the number defensible"],
 ["5. Reason line","Feature attribution, not free text","Top contributing features (APL variable contributions / SHAP) mapped to human-readable templates — the agent always sees WHY (governance requirement, Section 4)"],
 ["6. Ranking","Not the % alone","Sort order = propensity x expected program value, so a 60% chance at a high-value DR enrollment can outrank a 90% chance at eBill; the displayed % remains the propensity component"],
], widths=[1.1,1.9,3.8], font_size=8)
bullets([
 ("Cold start","at go-live there is no outcome history: Phase 1 begins with rules-based heuristic scoring (weighted business rules, e.g. registered EV charger adds strongly to EV Managed Charging) presented identically in the UI, swapped for the trained model after ~3-6 months of outcomes."),
 ("Quality monitoring","calibration drift and recommendation acceptance rate reviewed quarterly (Section 5 KPIs); models retrained on a schedule, versioned in AI Launchpad."),
 ("One-sentence answer for stakeholders","'The percentage is a calibrated enrollment probability learned from our own enrollment outcomes, with the driving factors shown to the agent — in the demo it is illustrative seed data.'"),
])

doc.add_heading("4. Governance: ML Recommends, BRFplus Decides", level=1)
bullets([
 ("Decision boundary","hard eligibility (E), warnings (W) and pre-active behavior remain exclusively BRFplus decision-table outcomes — auditable, transportable, regulator-facing. No model score can enroll, reject or override; a regulator asking 'why was this medical customer denied PrePay?' always gets message 052 from a decision table."),
 ("AI operating zone","ranking, ordering, explaining, predicting, drafting and extracting — always with the agent in the loop and always labeled as AI-generated in the UI."),
 ("Explainability","recommendation reasons shown to the agent (feature-based, per Section 3 step 5); generative outputs grounded on the catalog with source attribution; no free hallucination surface."),
 ("Data protection","features minimized to identifiers + behavioral aggregates; PII never leaves the BTP subaccount boundary; generative prompts contain no bank/payment data; retention aligned with the ILM policy (CS-02)."),
 ("Model lifecycle","AI Launchpad for versioning/monitoring; drift alerts; quarterly business review of recommendation acceptance rate alongside the BRFplus governance meeting."),
 ("Bias safeguard","assistance-program recommendations (CARE/LIHEAP/PIPP) reviewed for demographic bias before go-live; recommendation may only ADD assistance offers, never suppress them."),
])

# ===== 4 ARCHITECTURE =====
doc.add_heading("5. Architecture Integration (no redesign required)", level=1)
table(["Component","Addition","Fits Where"], [
 ["EC-10 Scoring API (new)","POST /enrollment/v1/recommendations:score - returns per-program propensity + reason codes","New CAP endpoint beside EC-01..09; called by the UI after EC-02; scores cached per interaction"],
 ["Feature layer","Aggregated features (usage, payment behavior, participation) refreshed daily","HANA Cloud (existing instance) — calculation views over EnrollmentRequest/history + replicated FI-CA aggregates; SAP Datasphere if the client standardizes there"],
 ["Model runtime","Classification/ranking models","HANA Cloud PAL/APL in-database (Phase 1, zero new infrastructure) → SAP AI Core when MLOps maturity requires it"],
 ["Generative layer","Grounded explanation / copilot skills","Generative AI Hub (AI Core) with catalog retrieval; surfaced in the EC UI panel and optionally as a Joule Studio skill in SC V2"],
 ["Document AI","Income-proof extraction pipeline","SAP Document Information Extraction called from the EC-08 document flow"],
 ["Telemetry","Recommendation shown/accepted/ignored events","AuditLog extension (action = AI_RECO_*) — closes the learning loop and measures lift"],
], widths=[1.6,2.6,2.6], font_size=8.5)

# ===== 5 PHASING =====
doc.add_heading("6. Phasing and KPIs", level=1)
table(["Phase","Scope","Entry Condition","Success KPI"], [
 ["AI-Phase 1 (with Increment 2)","AI-1 recommendation + AI-2 grounded explanations; CX AI Toolkit call summaries","Enrollment Center live; 3+ months of history OR cold-start rules-based ranking","Recommendation acceptance rate >25%; program-per-call uplift; AHT reduction on program calls"],
 ["AI-Phase 2","AI-3 bill-impact advisor; AI-4 default risk; AI-5 document extraction","Interval data access via EDM/MDMS; document volume justifies automation","Rate-program regret rate <5%; DPA default reduction; verification backlog days"],
 ["AI-Phase 3","AI-6 exception triage; AI-7 DER forecasting (DER platform)","12+ months of exception/override labels; DER platform live","Exception MTTR; PTR settlement accuracy"],
], widths=[1.4,2.5,1.7,1.9], font_size=8.5)
p("Demo note: the interactive demo includes a working preview of AI-1 and AI-2 (mock scores and grounded "
  "explanation composed from catalog + customer facts) so stakeholders can see the agent experience today — "
  "clearly labeled as AI and kept outside the eligibility decision path, exactly like the production design.")

out = r"C:\Users\jnamm\OneDrive\Desktop\Service Cloud for Utilities\Enrollment Center\Enrollment Center - AI-ML Roadmap v1.0.docx"
doc.save(out)
print("Saved:", out)
