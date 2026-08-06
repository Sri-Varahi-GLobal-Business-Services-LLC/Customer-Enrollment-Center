# -*- coding: utf-8 -*-
"""Generate Enrollment Center Technical Design Document: S/4HANA data sources, APIs, CDS views, BRFplus, BTP architecture."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = "1F4E79"; HEADER_FILL = "1F4E79"; ALT_FILL = "DEEAF6"
doc = Document()
normal = doc.styles["Normal"]; normal.font.name = "Arial"; normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
for lvl, sz in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
    st = doc.styles[lvl]; st.font.name = "Arial"; st.font.size = Pt(sz); st.font.bold = True
    st.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    rpr = st.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None: rf = OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a), 'Arial')

def p(text, bold=False, size=None, color=None, align=None):
    para = doc.add_paragraph(); run = para.add_run(text); run.bold = bold
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor.from_string(color)
    if align: para.alignment = align
    return para

def bullets(items):
    for it in items:
        if isinstance(it, tuple):
            para = doc.add_paragraph(style="List Bullet")
            r = para.add_run(it[0] + " — "); r.bold = True
            para.add_run(it[1])
        else:
            doc.add_paragraph(it, style="List Bullet")

def numbered(items):
    for it in items: doc.add_paragraph(it, style="List Number")

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def table(headers, rows, widths=None, font_size=9.5):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(font_size)
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); shade(c, HEADER_FILL)
    for r_i, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val)); run.font.size = Pt(font_size)
            if r_i % 2 == 1: shade(cells[i], ALT_FILL)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ============ COVER ============
for _ in range(6): doc.add_paragraph()
p("Customer Enrollment Center on SAP BTP", bold=True, size=26, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER)
p("Technical Design Document", bold=True, size=18, color="404040", align=WD_ALIGN_PARAGRAPH.CENTER)
p("S/4HANA Utilities Data Sources, Standard APIs, CDS Views, BRFplus Applications and End-to-End Architecture",
  size=12, color="595959", align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(7): doc.add_paragraph()
table(["Attribute","Value"], [
 ["Document","Enrollment Center – Technical Design (TDD)"],
 ["Version / Status","1.1 (program catalog extended from the U.S. Utility Customer Programs Reference)"],
 ["Date","08 July 2026"],
 ["Related Documents","Enrollment Center – SAP Service Cloud V2 Solution Design v2.0; CS-01 / CS-02 BPDs; U.S. Utility Customer Programs Reference Catalog (Aug 2026); interactive demo (3 URLs)"],
 ["Grounding Sources","SAP S/4HANA Utilities CDS view catalog and API guide (client reference library); legacy CRM Enrollment Center FS/TS (FUNC_PGMENRL_PRECHECK, RS_DPP_INT_VALIDATION, ZRFC_GET_RATE_CATEGORY); U.S. Utility Customer Programs Reference (program groups, SAP representation model)"],
 ["Audience","BTP developers, S/4HANA Utilities (ISU/FI-CA) developers, integration architects, Basis/security"],
], widths=[1.7,4.8])
doc.add_page_break()

# ============ 1 SCOPE ============
doc.add_heading("1. Purpose and Scope", level=1)
p("This TDD specifies the technical building blocks for the custom Enrollment Center application: the BTP "
  "application components, the S/4HANA Utilities data sources it reads (CDS views, released APIs, underlying "
  "tables), the BRFplus applications that decide program eligibility, and the custom objects that bridge the "
  "two. It is the developer-facing companion to the Solution Design v2.0: functional behavior (four UI panels, "
  "dual access points, deferred move-in fulfillment, status model) is defined there and not repeated here.")
p("Design rule restated: S/4HANA Utilities is the system of record; the BTP application persists only process "
  "state (enrollment requests, audit) and reads all master/transactional context live from S/4HANA. BRFplus in "
  "S/4HANA remains the single eligibility rules engine, callable by every channel.")

# ============ 2 ARCHITECTURE ============
doc.add_heading("2. End-to-End Architecture", level=1)
doc.add_heading("2.1 Runtime Components", level=2)
table(["Layer","Component","Technology","Responsibility"], [
 ["SC V2","Agent desktop + mashup / Move-In step","SAP Service Cloud V2 UI integration","Launch Enrollment Center with signed context (BP, CA, case / move-in ID)"],
 ["BTP – UI","Enrollment Center UI","SAPUI5 (TypeScript, Horizon), HTML5 App Repository, managed approuter","Four-panel UI, enroll/de-enroll wizards, mode switch (standalone / move-in)"],
 ["BTP – Service","Enrollment Service","CAP (Node.js) on Cloud Foundry","APIs /enrollment/v1/*; state machine; validation; audit; program catalog"],
 ["BTP – Rules facade","Eligibility Adapter","CAP module + destination (RFC/OData)","Batches eligibility calls to S/4HANA BRFplus wrapper; normalizes messages/severities"],
 ["BTP – Persistence","Enrollment DB","SAP HANA Cloud (HDI)","EnrollmentRequest, ProgramCatalog, OverrideRecord, AuditLog, EventInbox"],
 ["BTP – Workflow","Exception & override workflow","SAP Build Process Automation","Supervisor override, Error-In-Review handling"],
 ["Integration","iFlows IF-01…IF-09","SAP Integration Suite (Cloud Integration + API Management)","Mediation to S/4HANA OData/RFC/SOAP; mapping; retry; DLQ"],
 ["Eventing","Event backbone","SAP Event Mesh","utilco/s4/movein/*, utilco/s4/enrollment/*, utilco/ec/fulfillment/*"],
 ["Connectivity","Destinations + Cloud Connector","BTP Destination/Connectivity service","Principal propagation (OAuth2SAMLBearer) and RFC tunnel to on-premise S/4HANA"],
 ["Identity","IAS + XSUAA","SAP Cloud Identity Services","SSO from SC V2; role collections EC_Agent / EC_Supervisor / EC_Admin"],
 ["Backend","S/4HANA Utilities","ISU + FI-CA + BRFplus + BPEM","Master/transactional data, rules, fulfillment objects, exception cases"],
], widths=[0.9,1.5,1.7,2.4], font_size=8.5)
doc.add_heading("2.2 Program Grouping and System-of-Record Model", level=2)
p("Per the U.S. Utility Customer Programs Reference, programs are grouped by the dimension of the customer "
  "relationship they change — that dimension determines the SAP object that carries them and therefore the "
  "fulfillment recipe. The Enrollment Center is the single agent front door for all groups; ownership differs:")
table(["Group","Programs (examples)","SAP Representation / Owner","Recipe Family"], [
 ["Rate / pricing","TOU, CPP, Dynamic Peak Pricing (DPP — the legacy Enrollment Center's rate program, fulfilled as product change to the DPP/RCPP rate), PTR rate variants, EV rate, Green Power, Community Solar, NEM, Interruptible tariff","Rate category / utilities product, product options and surcharge operands on the installation (IS-U)","RECIPE_RATE"],
 ["Bill-calculation behavior","Budget Billing, Summary/Consolidated Billing","Budget billing plan; joint invoicing / bill grouping (IS-U billing config)","RECIPE_BB, RECIPE_BILLCONFIG"],
 ["Payment behavior","AutoPay, CSDD, Deferred Payment Arrangement (DPA), Payment Extension, PrePay, AMP","Contract account (FI-CA): payment method, due-date terms, installment plans, deferrals, prepay settings","RECIPE_PAYMENT, RECIPE_DPA, RECIPE_PREPAY, RECIPE_CSDD"],
 ["Communication / channel","eBill, Usage/High-Bill Alerts, Third-Party Notification","Correspondence control + BP communication data; notification platform for alert delivery","RECIPE_EBILL, RECIPE_CORR, RECIPE_NOTIF"],
 ["Assistance / protected status","CARE/FERA, LIHEAP, PIPP, Medical Baseline / Life Support","BP/CA status marks + rate discount + dunning protection (IS-U / FI-CA)","RECIPE_ASSIST"],
 ["DER / DR lifecycle","Smart Thermostat/BYOD, DLC/Peak Power Savers, PTR events, EV Managed Charging, battery/VPP, C&I curtailment, OpenADR, EE rebates","NOT native IS-U: full lifecycle (eligibility, enrollment, events, performance, incentives) owned by the BTP DER platform; only a thin date-effective participation reference in S/4HANA","RECIPE_DER, RECIPE_DEVICE, RECIPE_EE"],
], widths=[1.1,1.9,2.2,1.3], font_size=8)
p("Modeling rule (key decision from the reference): groups 1–5 are configuration on existing IS-U objects and "
  "are fulfilled by the Section 5.3 wrappers; group 6 is a true program lifecycle — the Enrollment Service "
  "delegates it to the DER platform APIs and stores only the cross-reference. The configurable "
  "Program/ProgramVersion catalog construct applies to group 6; groups 1–5 map to fixed recipe types.")
doc.add_heading("2.3 Data Flow Summary", level=2)
numbered([
 "UI loads → Enrollment Service GET /context aggregates: BP (API_BUSINESS_PARTNER), premise/installation/contract (utilities CDS OData), contract account + financial flags (FI-CA services), device/AMI capability (device CDS), participation (Z-table OData + HANA Cloud state).",
 "POST /eligibility:evaluate → one RFC call to Z_EC_ELIGIBILITY_PRECHECK → BRFplus evaluates all program rulesets for all contract accounts in parallel → messages (severity E/W/I) returned and cached for the interaction.",
 "POST /enrollments → server-side re-validation → persist (HANA Cloud) → fulfillment recipe dispatch via Cloud Integration to the S/4HANA API for the program type (rate change, installment plan, account attribute, service order, prepay).",
 "S/4HANA emits confirmation/business events → Event Mesh → Enrollment Service updates status, writes SC V2 case timeline.",
 "Move-in mode: requests parked; utilco/s4/movein/completed releases them (re-validate + dispatch); cancellation cascades.",
])

# ============ 3 DATA REQUIREMENTS ============
doc.add_page_break()
doc.add_heading("3. Data Requirements and S/4HANA Sources", level=1)
p("Mapping of every data need of the Enrollment Center to its S/4HANA Utilities source. CDS view names are "
  "taken from the S/4HANA Utilities view catalog in the client reference library; underlying tables are listed "
  "for developer orientation and debugging — all access goes through CDS/APIs, never direct table reads.")
table(["Data Need (UI element)","CDS View / API Source","Underlying Tables (reference)"], [
 ["Business partner header, addresses, comm data","API_BUSINESS_PARTNER (OData V2, A2X)","BUT000, BUT020/ADRC, BUT0BK"],
 ["Contract account attributes, locks, dunning level","FI-CA contract account read (custom OData Z_CA_SRV wrapping FKK_* reads; SOAP CONTRACTACCOUNTBULKREPLICATEREQUEST_OUT exists for replication scenarios)","FKKVK, FKKVKP, DFKKLOCKS"],
 ["Open items / arrears (DPP qualification)","Custom CDS Z_C_CaOpenItems on FI-CA line items","DFKKOP, DFKKKO"],
 ["Installment plan existence (DPA hard stop)","FI-CA installment plan check inside BRFplus rules (FM-based)","DFKKOP (statistical items), installment plan header via FI-CA FMs"],
 ["Budget billing participation","Custom CDS / FM check in rules","EABP (budget billing plan)"],
 ["Premise / service location","I_UtilitiesPremise (+ I_UtilsPremiseKeyDateAnalysis)","EVBS"],
 ["Connection object","I_UtilitiesConnectionObject","IFLOT/ILOA (functional location), EHAU"],
 ["Installation incl. rate category (product change basis)","I_UtilitiesInstallation, I_UtilitiesInstallationHist","EANL, EANLH (TARIFTYP = rate category)"],
 ["Utilities contract (active/history, move-in date)","I_UtilitiesContract, I_UtilitiesContractHistory, I_UtilsContractChangeHistory","EVER"],
 ["Move-out / move-in context","I_UtilitiesMoveOutDocument + move-in APIs (Section 5)","EVER/EANLH change docs"],
 ["Point of delivery","I_UtilsAllocOfInstToPoD, external PoD mapping","EUIHEAD, EUITRANS, EUIINSTLN"],
 ["Installed device + AMI capability (PTR/CPP/PrePay checks)","I_UtilsInstalledDevice, I_UtilitiesDeviceHeader/History, I_UtilsAdvncdMeteringSystem, C_UtilsInstalledDeviceQuery","EQUI, EQKT, EGERH, ETDZ"],
 ["Meter reading context (validation prompts)","I_MeterReadingDocument, C_MeterReadingDocumentQuery","EABL, EABLG"],
 ["Billing history (eligibility rules: 12-month history, high-bill context)","I_ActualUtilsBillingDocument, C_UtilsContractToBeBilledQ","ERCH, ERDK, DBERCHZ*"],
 ["Owner allocation (landlord scenarios)","I_UtilitiesOwnerAllocation, I_AllocOfInstToOwnership","EVBS/EIGENTUM structures"],
 ["BPEM exception cases (status panel / exception links)","I_BPEMCaseQuery / C_BPEMCaseQuery","EMMA_CASE"],
 ["Program participation + enrollment history","Custom Z-table OData Z_EC_ENROLLMENT_SRV (Section 6) + HANA Cloud process store","ZEC_ENROLL_HDR, ZEC_ENROLL_HIST (custom)"],
 ["Rate catalog / allowed target rates","Re-use legacy services ZRFC_GET_RATE_CATEGORY, Z_MIMO_MI/MO_RES_ALLOWED_RATES","Rate config tables via FM"],
], widths=[1.9,2.5,2.1], font_size=8.5)
p("Note: table names are for orientation and debugging; where a released CDS view exists it is the mandated "
  "interface. Objects marked custom (Z…) are specified in Section 6. Names of FI-CA function-module-based "
  "checks are finalized during realization against the installed release.")

# ============ 4 CDS LIST ============
doc.add_heading("4. CDS View Inventory (from the S/4HANA Utilities catalog)", level=1)
table(["Area","CDS Views"], [
 ["Master data","I_UtilitiesPremise, I_UtilitiesConnectionObject, I_UtilitiesInstallation, I_UtilitiesInstallationHist, I_UtilitiesContract, I_UtilitiesContractHistory, I_UtilitiesSalesContract, I_UtilitiesSalesContractItem, I_UtilitiesOwnerAllocation, I_AllocOfInstToOwnership, I_UtilitiesServiceProvider"],
 ["Device / metering","I_UtilitiesDeviceHeader, I_UtilitiesDeviceHistory, I_UtilsInstalledDevice, I_UtilsTechlyInstalledDevice, I_UtilsDeviceCategory, I_UtilsAdvncdMeteringSystem (AMI), I_UtilitiesRegister, I_UtilsDeviceRateData, C_UtilsInstalledDeviceQuery, C_UtilsDeviceStockQuery"],
 ["Meter reading","I_MeterReadingDocument, I_MeterReadingDocumentReason, I_MeterReadingUnit, C_MeterReadingDocumentQuery, C_MissingMtrRdngOrderQuery"],
 ["Billing / invoicing","I_ActualUtilsBillingDocument, I_ReversedUtilsBillingDocument, I_UtilsBillgOrderForContract, I_UtilsContractToBeBilled, C_ActualUtilsBillingDocumentQ, C_UtilitiesInvoicingDocumentQ, C_UtilsContractBlkdForBillgQ"],
 ["Analysis / key-date","I_UtilsContractKeyDateAnalysis, I_UtilsInstKeyDateAnalysisC, I_UtilsPremiseKeyDateAnalysis, I_UtilsContractChangeHistory, C_UtilsContrKeyDateAnalysisQ"],
 ["Exception mgmt","I_BPEMCaseQuery, I_BPEMCaseCube, C_BPEMCaseQuery"],
 ["Move processes","I_UtilitiesMoveOutDocument"],
 ["Data export (federation to marketing/CX)","C_UtilsContractHistoryDEX, C_UtilsInstallationHistoryDEX"],
], widths=[1.3,5.2], font_size=9)
p("Consumption pattern: expose the required I_* views as OData services via custom service bindings "
  "(ABAP RESTful service binding or SEGW gateway projects per release), grouped into one consolidated "
  "service Z_EC_CONTEXT_SRV to minimize round trips from Cloud Integration (IF-01 performs parallel reads).")

# ============ 5 API LIST ============
doc.add_page_break()
doc.add_heading("5. Standard API Inventory", level=1)
doc.add_heading("5.1 OData APIs", level=2)
table(["API","Type","Use in Enrollment Center"], [
 ["API_BUSINESS_PARTNER","OData V2 (A2X)","BP read: names, addresses, communication, bank (display-only)"],
 ["Utilities CDS-based OData (Section 4)","OData V2/V4 service bindings","Premise, installation, contract, device, billing context reads"],
 ["BPEM case OData (on I_BPEMCaseQuery)","OData","Exception case status linkage"],
], widths=[2.3,1.3,2.9], font_size=9)
doc.add_heading("5.2 SOAP Services (from the client's S/4HANA Utilities API guide)", level=2)
table(["Service (technical name)","Direction","Use"], [
 ["CONTRACTACCOUNTBULKREPLICATEREQUEST_OUT","S/4 → out","Contract account replication (used by SC V2 utilities federation; EC reads replicated key data where present)"],
 ["UtilitiesPremiseBulkReplicateRequest_Out","S/4 → out","Premise replication to CX/marketing consumers"],
 ["UtilitiesContractBulkReplicateRequest_Out","S/4 → out","Utilities contract replication"],
 ["UtilitiesInstallationBulkReplicateRequest_Out","S/4 → out","Installation + installation facts replication"],
 ["UtilitiesSalesContractBulkReplicateRequest_Out","S/4 → out","Sales contract replication"],
 ["UtilitiesGeneralDataBulkRequest_Out","S/4 → out","General utilities data bulk requests"],
], widths=[2.9,0.9,2.7], font_size=9)
p("These A2A services support the surrounding landscape (SC V2 utilities add-on, marketing). The Enrollment "
  "Center itself prefers synchronous CDS/OData reads; the replication services are listed because IF-01 falls "
  "back to replicated data for latency-critical fields where the landscape already maintains them.")
doc.add_heading("5.3 RFC / BAPI Layer (fulfillment writes)", level=2)
table(["Operation","Interface (representative — verify against release)","Program Types"], [
 ["Move-in creation (from SC V2 guided flow)","Standard utilities move-in API / ISU move-in FM wrapped as Z_EC_MOVEIN_SRV","Move-in embedded mode"],
 ["Rate/product change on installation","Installation change FMs (EANLH/TARIFTYP update) wrapped as Z_EC_RATECHANGE; re-use legacy Z_GET_NEW_MI_RATE_CAT, Z_MIMO_* allowed-rate checks","DPP (Dynamic Peak Pricing), PTR, CPP, TOU, CARE rate variants, Green Power, Community Solar"],
 ["Installment plan create/deactivate","FI-CA installment plan FMs wrapped as Z_EC_INSTPLAN","DPA (Deferred Payment Arrangement)"],
 ["Budget billing plan create/end","Budget billing FMs wrapped as Z_EC_BUDGETBILL","Budget Billing"],
 ["Contract account attribute update (due date, eBill, prepay flag)","Contract account change FM/BAPI wrapped as Z_EC_CAUPDATE","CSDD, eBill, PrePay settings"],
 ["Service order creation (device programs)","Service order API wrapped as Z_EC_SERVICEORDER","Smart Thermostat, PPS, AMI exchange"],
 ["Payment method / extension / deferral","FI-CA payment method update + due-date deferral FMs wrapped as Z_EC_PAYMENT","AutoPay, Payment Extension"],
 ["Assistance status + discount","BP/CA status marks + rate discount operand wrapped as Z_EC_ASSIST; AMP forgiveness schedule via FI-CA","LIHEAP, PIPP, AMP, CARE/FERA, Medical Baseline"],
 ["Correspondence recipient / bill copy","Correspondence control update wrapped as Z_EC_CORR","Third-Party Notification, eBill"],
 ["Joint invoicing / bill grouping","Bill grouping config wrapped as Z_EC_BILLGROUP","Summary / Consolidated Billing"],
 ["Product option / surcharge operand","Installation operand update via Z_EC_RATECHANGE","Green Power, Community Solar"],
 ["DER platform handoff","DER platform enrollment API (BTP-to-BTP, no S/4HANA write except participation reference)","EV Managed Charging, OpenADR, thermostat/DLC dispatch programs"],
 ["Alert subscription","Notification platform API (no S/4HANA write)","Usage / High-Bill Alerts"],
 ["EE rebate intake","EE/rebate system API or S/4HANA service order per landscape","Energy-Efficiency Rebates"],
 ["Eligibility pre-check","Z_EC_ELIGIBILITY_PRECHECK (successor of legacy FUNC_PGMENRL_PRECHECK)","All programs"],
 ["De-enrollment exit check","Z_EC_DEENROLL_CHECK","All programs"],
], widths=[1.9,3.0,1.6], font_size=8.5)
doc.add_heading("5.4 Events (Event Mesh topics)", level=2)
table(["Topic","Producer","Payload / Trigger"], [
 ["utilco/s4/movein/completed","S/4HANA (event enablement on move-in completion)","Move-in ID, BP, CA, contract, activation date — releases parked Enrollment Requests"],
 ["utilco/s4/movein/cancelled","S/4HANA","Move-in ID — cascade-cancels linked requests"],
 ["utilco/s4/enrollment/confirmed","S/4HANA outbound event from fulfillment commit","Request ID, backend object keys, actual dates"],
 ["utilco/s4/bpem/casechanged","S/4HANA BPEM","Case number/status — updates Error-In-Review items"],
 ["utilco/ec/fulfillment/dispatch","Enrollment Service","Recipe execution command consumed by IF-03"],
], widths=[2.0,1.7,2.8], font_size=9)

# ============ 6 BRFPLUS ============
doc.add_page_break()
doc.add_heading("6. BRFplus Application Design", level=1)
doc.add_heading("6.1 Application Structure", level=2)
p("One BRFplus application ZEC_PROGRAM_ELIGIBILITY in S/4HANA (client-independent storage, transportable), "
  "modeled on the proven legacy design (DPP/DLA/PrePay rulesets called through one pre-check function). The "
  "legacy ruleset RS_DPP_INT_VALIDATION and function FUNC_PGMENRL_PRECHECK are migrated/renamed into this "
  "application.")
table(["BRFplus Object","Name","Purpose"], [
 ["Application","ZEC_PROGRAM_ELIGIBILITY","Container; storage type S (system); transported via workbench request"],
 ["Function","F_EC_ELIGIBILITY_ALL","Entry point: evaluates all requested programs for all supplied contract accounts; mode PRE_ACTIVE for move-in context"],
 ["Function","F_EC_DEENROLL_CHECK","Exit-rule evaluation per participation"],
 ["Ruleset (per program)","Rate/DR: RS_DPP_INT_VALIDATION (migrated legacy ruleset — Dynamic Peak Pricing), RS_PTR, RS_CPP, RS_TOU, RS_PPS, RS_THERM, RS_EVMC, RS_INTR, RS_ADR, RS_CDR, RS_GREEN, RS_CSOL. Billing/payment: RS_CSDD, RS_EBILL, RS_DPA (installment plan), RS_BB, RS_PREPAY, RS_AUTOPAY, RS_PEXT, RS_SUMBILL. Assistance/account: RS_CARE, RS_LIHEAP, RS_PIPP, RS_AMP, RS_MED, RS_TPN, RS_ALERTS, RS_EEREB (suffix _VALIDATION on all)","All rules of a ruleset execute completely — every failure is returned (no first-fail exit), per legacy design. Note: legacy 'DPP' = Dynamic Peak Pricing (rate program); the payment plan is DPA"],
 ["Decision tables","DT_<PGM>_HARDSTOPS, DT_<PGM>_WARNINGS, DT_WAITING_PERIODS, DT_RATECLASS_MATRIX, DT_PROGRAM_EXCLUSIONS","Maintainable rule content: rate class restrictions, AMI requirement, waiting periods, mutual program exclusions (e.g., PrePay vs Medical Alert), credit thresholds"],
 ["Expressions","EX_AMI_CHECK, EX_ARREARS_AMOUNT, EX_RETURNED_PAYMENTS_12M, EX_ACTIVE_INSTPLAN, EX_MEDICAL_FLAG, EX_BILLING_HISTORY_MONTHS","DB lookups/procedure calls feeding the decision tables"],
 ["Context data object","CTX_EC_ELIG (structure)","BP, CA list, division, rate category, premise, installation, device AMI flag, dunning level, medical flag, program history table, mode (STANDALONE/PRE_ACTIVE)"],
 ["Result data object","RES_EC_ELIG (table)","Program ID, CA, eligible flag, message class/number, severity (E=hard stop, W=warning, I=not-yet-eligible), message variables"],
], widths=[1.3,2.6,2.6], font_size=8.5)
doc.add_heading("6.2 Invocation and Performance", level=2)
bullets([
 "Wrapper RFC Z_EC_ELIGIBILITY_PRECHECK builds CTX_EC_ELIG (parallel CDS reads via aRFC), calls F_EC_ELIGIBILITY_ALL once for all programs and contract accounts, and returns RES_EC_ELIG — one round trip per Enrollment Center load, matching the legacy contract.",
 "BRFplus functions generated (not interpreted) for performance; generation refreshed in transport pipeline.",
 "Rule content changes (decision table rows) are business-maintainable via BRFplus workbench with change-log; structural changes follow WRICEF governance.",
 "PRE_ACTIVE mode: rules referencing billing history or active contract return severity I ('not yet eligible') instead of E, enabling the move-in catalog behavior.",
 "All messages come from message class ZEC_ELIG (number ranges: 0xx hard stops, 3xx warnings, 9xx not-yet-eligible) so the UI and audit log are language-enabled.",
])

# ============ 7 CUSTOM OBJECTS ============
doc.add_heading("7. Custom Development Inventory (S/4HANA side)", level=1)
table(["Object","Type","Description"], [
 ["Z_EC_ELIGIBILITY_PRECHECK","RFC FM","BRFplus wrapper; exposed to BTP via destination (RFC) or thin OData"],
 ["Z_EC_DEENROLL_CHECK","RFC FM","Exit rules: minimum participation, balances, pending device orders"],
 ["Z_EC_CONTEXT_SRV","OData service","Consolidated context read over the Section 4 CDS views"],
 ["Z_EC_ENROLLMENT_SRV","OData service","CRUD on ZEC_ENROLL_HDR/ZEC_ENROLL_HIST program participation tables"],
 ["ZEC_ENROLL_HDR / ZEC_ENROLL_HIST","Tables","Program participation + history in S/4HANA (system of record for participation; BTP holds process state only)"],
 ["Z_EC_RATECHANGE / Z_EC_INSTPLAN / Z_EC_BUDGETBILL / Z_EC_CAUPDATE / Z_EC_SERVICEORDER / Z_EC_MOVEIN_SRV","RFC/OData wrappers","Fulfillment recipe executors (Section 5.3)"],
 ["Outbound event enablement","Event config + Z-exit where needed","Topics in Section 5.4 via SAP event enablement framework"],
 ["ZEC_ELIG message class","Message class","Eligibility/exit messages with severity convention"],
 ["BPEM case categories ZEC1/ZEC2","BPEM config","Enrollment fulfillment errors; auto-close event on resolution"],
], widths=[2.1,1.2,3.2], font_size=9)

# ============ 8 BTP APP ============
doc.add_page_break()
doc.add_heading("8. BTP Application Design", level=1)
doc.add_heading("8.1 CAP Service Model", level=2)
table(["Entity / Service","Key Fields","Notes"], [
 ["EnrollmentRequest","requestId (UUID), programId, businessPartner, contractAccount, moveInId?, status, requestDate, actualDate, payload (JSON), idempotencyKey","Status machine per Solution Design 6.4; unique index on idempotencyKey"],
 ["ProgramCatalog","programId, version, name, category, recipeType, captureForm (JSON), active","Configuration-over-code: new program = catalog row + BRFplus ruleset"],
 ["OverrideRecord","requestId, messageCode, approver, justification, decidedAt","Warning overrides; immutable"],
 ["AuditLog","seq, requestId?, actor, action, before/after, hash","Append-only; exported to client SIEM"],
 ["EventInbox","eventId, topic, payload, processedAt, status","Idempotent event consumption + reconciliation source"],
 ["Service enrollment/v1","context, programs, eligibility:evaluate, enrollments, deenrollments, exceptions","OpenAPI published in API Management; consumed by UI and future channels"],
], widths=[2.0,2.6,1.9], font_size=8.5)
doc.add_heading("8.2 Destinations", level=2)
table(["Destination","Target","Auth"], [
 ["S4_UTILITIES_ODATA","S/4HANA gateway (Z_EC_CONTEXT_SRV, API_BUSINESS_PARTNER, CDS services)","OAuth2SAMLBearerAssertion (principal propagation) via Cloud Connector"],
 ["S4_UTILITIES_RFC","Z_EC_ELIGIBILITY_PRECHECK, fulfillment wrappers","RFC via Cloud Connector, technical user + agent ID in payload"],
 ["SCV2_API","SC V2 REST (case timeline)","OAuth2ClientCredentials"],
 ["EVENT_MESH","Event Mesh instance","Service key (OAuth2)"],
], widths=[1.7,3.1,1.7], font_size=9)
doc.add_heading("8.3 Non-Functionals", level=2)
bullets([
 "Context + eligibility render target < 3 s: parallel destination calls, 60-s TTL cache on static catalog, single-round-trip eligibility.",
 "Idempotent writes end-to-end (idempotencyKey on POST /enrollments; EventInbox de-duplication).",
 "Reconciliation job (CAP scheduled task) compares Pending Backend items with S/4HANA state daily.",
 "Observability: BTP Alert Notification + Cloud Logging; correlation ID propagated SC V2 → BTP → iFlow → S/4HANA (passport).",
 "Transports: CAP/UI via CI/CD (BTP Continuous Integration & Delivery) + Transport Management; ABAP objects via CTS; BRFplus via workbench transport with generation step.",
])

# ============ 9 SECURITY ============
doc.add_heading("9. Security Summary", level=1)
bullets([
 "IAS is the IdP for SC V2 and BTP; XSUAA role collections EC_Agent (read+enroll+de-enroll), EC_Supervisor (+override, exceptions), EC_Admin (catalog).",
 "Principal propagation to S/4HANA for OData reads (named-user authorization applies); RFC fulfillment uses a least-privilege technical user with agent ID logged.",
 "S/4HANA authorization objects checked in wrappers (e.g., contract account change, installment plan create) — BTP roles never widen backend authority.",
 "Audit: AuditLog (BTP) + change documents on ZEC_ENROLL_* + BRFplus change log + BTP Audit Log service.",
 "No card/bank capture in the Enrollment Center; personal data minimized to identifiers; retention aligned with S/4HANA ILM.",
])

# ============ 10 OPEN POINTS ============
doc.add_heading("10. Open Points for Realization", level=1)
numbered([
 "Confirm installed S/4HANA release and verify FI-CA function-module names for installment plan / budget billing wrappers (Section 5.3 marked representative).",
 "Confirm event enablement approach for utilities move-in completion on the installed release (standard event vs. Z-exit publisher).",
 "Decide RFC vs. OData exposure of Z_EC_ELIGIBILITY_PRECHECK (RFC preferred for payload size; OData acceptable if Cloud Connector RFC is restricted).",
 "Agree participation system of record detail: ZEC_ENROLL_* in S/4HANA (recommended, stated above) vs. program-specific standard objects only.",
 "Sizing: HANA Cloud (starter), Cloud Integration tier, Event Mesh consumption estimate based on enrollment volumes.",
])

out = r"C:\Users\jnamm\OneDrive\Desktop\Service Cloud for Utilities\Enrollment Center\Enrollment Center - Technical Design (S4HANA Data Sources, APIs, CDS, BRFplus) v1.0.docx"
doc.save(out)
print("Saved:", out)
