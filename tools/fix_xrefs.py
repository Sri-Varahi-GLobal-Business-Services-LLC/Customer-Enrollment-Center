# -*- coding: utf-8 -*-
"""Fix stale section cross-references in SDD v2.0 after user's renumbering."""
from docx import Document

PATH = r"C:\Users\jnamm\OneDrive\Desktop\Service Cloud for Utilities\Enrollment Center\Enrollment Center - SAP Service Cloud V2 Solution Design v2.0.docx"
doc = Document(PATH)

REPLACEMENTS = [
    ("Section 5.3", "Section 6.3"),
    ("(see Section 9)", "(see Section 10)"),
]

count = 0
def fix_runs(para):
    global count
    for run in para.runs:
        for old, new in REPLACEMENTS:
            if old in run.text:
                run.text = run.text.replace(old, new)
                count += 1

for p in doc.paragraphs:
    fix_runs(p)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            for p in c.paragraphs:
                fix_runs(p)

doc.save(PATH)
print("Replacements made:", count)
