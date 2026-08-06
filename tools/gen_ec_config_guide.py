# -*- coding: utf-8 -*-
"""Generate Enrollment Center BTP Deployment, Configuration and S/4HANA Integration Guide."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT="1F4E79"; HEADER_FILL="1F4E79"; ALT_FILL="DEEAF6"
doc = Document()
n = doc.styles["Normal"]; n.font.name="Arial"; n.font.size=Pt(10.5); n.paragraph_format.space_after=Pt(6)
for lvl, sz in [("Heading 1",16),("Heading 2",13),("Heading 3",11.5)]:
    st = doc.styles[lvl]; st.font.name="Arial"; st.font.size=Pt(sz); st.font.bold=True
    st.font.color.rgb = RGBColor(0x1F,0x4E,0x79)
    rpr = st.element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
    if rf is None: rf = OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a),'Arial')

def p(text, bold=False, size=None, color=None, align=None):
    para = doc.add_paragraph(); r = para.add_run(text); r.bold=bold
    if size: r.font.size=Pt(size)
    if color: r.font.color.rgb=RGBColor.from_string(color)
    if align: para.alignment=align
    return para

def mono(text):
    para = doc.add_paragraph(); r = para.add_run(text)
    r.font.name="Consolas"; r.font.size=Pt(8.5)
    para.paragraph_format.space_after=Pt(2)
    return para

def bullets(items):
    for it in items:
        if isinstance(it, tuple):
            para = doc.add_paragraph(style="List Bullet")
            r = para.add_run(it[0]+" — "); r.bold=True
            para.add_run(it[1])
        else:
            doc.add_paragraph(it, style="List Bullet")

def numbered(items):
    for it in items: doc.add_paragraph(it, style="List Number")

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def table(headers, rows, widths=None, font_size=9):
    t = doc.add_table(rows=1, cols=len(headers)); t.style="Table Grid"
    t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""
        r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(font_size)
        r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); shade(c,HEADER_FILL)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=""
            run=cells[i].paragraphs[0].add_run(str(val)); run.font.size=Pt(font_size)
            if ri%2==1: shade(cells[i],ALT_FILL)
    if widths:
        for i,w in enumerate(widths):
            for row in t.rows: row.cells[i].width=Inches(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(2)
    return t

# ========== COVER ==========
for _ in range(6): doc.add_paragraph()
p("Customer Enrollment Center on SAP BTP", bold=True, size=26, color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER)
p("Deployment, Configuration and S/4HANA Integration Guide", bold=True, size=18, color="404040", align=WD_ALIGN_PARAGRAPH.CENTER)
p("BTP account setup, service configuration, connectivity, SC V2 and S/4HANA Utilities configuration, and detailed S/4HANA development specifications",
  size=12, color="595959", align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(7): doc.add_paragraph()
table(["Attribute","Value"], [
 ["Document","Enrollment Center – BTP Deployment & Configuration Guide"],
 ["Version / Status","1.0 (Draft for Review)"],
 ["Date","08 July 2026"],
 ["Related Documents","Solution Design v2.0; Technical Design v1.1 (data sources, APIs, CDS, BRFplus); CS-01/CS-02 BPDs"],
 ["Audience","BTP administrators, Basis, S/4HANA developers, integration developers, security team"],
 ["Assumed Landscape","SAP S/4HANA Utilities (on-premise or private cloud) + SAP Service Cloud V2 + SAP BTP (Cloud Foundry); three-tier DEV/QA/PRD"],
], widths=[1.7,4.8])
doc.add_page_break()

# ========== 1 OVERVIEW ==========
doc.add_heading("1. Introduction and Prerequisites", level=1)
p("This guide provides the end-to-end setup instructions to deploy the Enrollment Center: BTP account and "
  "service configuration, security and connectivity, Integration Suite and Event Mesh setup, SC V2 embedding, "
  "the configuration required inside S/4HANA Utilities, and the detailed inventory of S/4HANA developments. "
  "Follow the deployment sequence in Section 11; every step references the detailed section that specifies it.")
p("Prerequisites:", bold=True)
bullets([
 "BTP global account with entitlements per Section 2.2; authorization to create subaccounts and service instances.",
 "SAP Cloud Identity Services (IAS) tenant, with SC V2 already using it for SSO.",
 "S/4HANA Utilities system reachable from the data center where Cloud Connector will run; ABAP development access (workbench + customizing requests).",
 "SC V2 tenant admin access (communication systems, mashup/external app configuration).",
 "GitHub repositories (application source) reachable from BTP CI/CD; current repos: nammi2011/Customer-Enrollment-Center (docs), application repo to be created at project start."])

# ========== 2 BTP ACCOUNT ==========
doc.add_heading("2. BTP Account and Service Setup", level=1)
doc.add_heading("2.1 Account Model", level=2)
table(["Level","Name (proposed)","Notes"], [
 ["Global account","<Client> Utilities","Existing or new"],
 ["Directory","CX Extensions","Groups SC V2-related extensions"],
 ["Subaccounts","ec-dev, ec-qa, ec-prd","One per tier; same region as SC V2 tenant (e.g., us10/us21); Cloud Foundry enabled, one org per subaccount, space: enrollment"],
], widths=[1.3,1.9,3.3])
doc.add_heading("2.2 Entitlements and Services", level=2)
table(["Service","Plan (baseline)","Purpose"], [
 ["Cloud Foundry Runtime","MEMORY 2–4 GB per tier","CAP Enrollment Service + approuter"],
 ["SAP HANA Cloud","hana (min 2 vCPU/16 GB) shared DEV/QA; dedicated PRD","Process persistence, catalog, audit"],
 ["SAP HANA Schemas & HDI Containers","hdi-shared","CAP database artifacts"],
 ["Authorization & Trust Management (XSUAA)","application + broker","App security, role collections"],
 ["Destination service","lite","Backend destinations"],
 ["Connectivity service","lite","On-premise access via Cloud Connector"],
 ["SAP Integration Suite","standard_edition","Cloud Integration + API Management (IF-01…IF-09)"],
 ["SAP Event Mesh","default","Async topics (Section 6)"],
 ["SAP Build Process Automation","standard","Supervisor override + exception workflow"],
 ["HTML5 Application Repository","app-host + app-runtime","SAPUI5 Enrollment Center UI"],
 ["SAP Build Work Zone, standard (optional)","standard","Launchpad hosting if UI also offered outside SC V2"],
 ["Continuous Integration & Delivery","default","Build/deploy pipelines from GitHub"],
 ["Cloud Transport Management","standard","Transport of MTAs DEV→QA→PRD"],
 ["Audit Log service","default","Platform audit retention"],
 ["Alert Notification","standard","Ops alerting"],
], widths=[2.1,1.5,2.9], font_size=8.5)
doc.add_heading("2.3 Application Deployment Artifacts", level=2)
bullets([
 ("MTA structure","mta.yaml with modules: ec-srv (CAP Node.js), ec-db (HDI), ec-ui (HTML5 app), ec-approuter (managed via destination to HTML5 repo); resources: xsuaa, destination, connectivity, hana, event-mesh, autoscaler (PRD)."),
 ("Build & deploy","mbt build → cf deploy per space, executed by CI/CD pipeline on merge to main (DEV) and release tags (QA/PRD via Transport Management)."),
 ("Configuration per tier","destination names identical across tiers; only URLs/credentials differ — no code changes between tiers."),
])

# ========== 3 SECURITY ==========
doc.add_heading("3. Identity and Security Configuration", level=1)
numbered([
 "Establish trust: in each subaccount, Security → Trust Configuration → add the IAS tenant (OpenID Connect). Disable default SAP ID service for business users.",
 "Create IAS user groups: EC_AGENT, EC_SUPERVISOR, EC_ADMIN; assign call center users (or map from corporate IdP groups).",
 "Deploy xs-security.json (Appendix A) with scopes ec.Agent, ec.Supervisor, ec.Admin and role templates; build role collections EnrollmentCenter_Agent / _Supervisor / _Admin in each subaccount and map them to the IAS groups.",
 "SC V2 SSO: SC V2 and the Enrollment Center UI share the IAS tenant, so the embedded UI authenticates silently in the agent session (no second login).",
 "Configure principal propagation trust for S/4HANA (Section 4.3).",
])

# ========== 4 CONNECTIVITY ==========
doc.add_page_break()
doc.add_heading("4. Connectivity: Cloud Connector and Destinations", level=1)
doc.add_heading("4.1 Cloud Connector Installation", level=2)
bullets([
 "Install two Cloud Connector instances (HA pair) in the client network zone with access to S/4HANA; connect each to ec-dev/ec-qa/ec-prd subaccounts with Location ID EC_MAIN.",
 "Expose S/4HANA as virtual hosts: s4-odata:44300 (HTTPS → internal gateway host) and s4-rfc (RFC → message server / application servers).",
])
doc.add_heading("4.2 Exposed Resources (least privilege)", level=2)
table(["Virtual Host","Protocol","Accessible Resources"], [
 ["s4-odata:44300","HTTPS","/sap/opu/odata/sap/API_BUSINESS_PARTNER/*; /sap/opu/odata/sap/Z_EC_CONTEXT_SRV/*; /sap/opu/odata/sap/Z_EC_ENROLLMENT_SRV/*; CDS service paths per TDD Section 4"],
 ["s4-rfc","RFC","Function name prefixes: Z_EC_* (exact-name entries for Z_EC_ELIGIBILITY_PRECHECK, Z_EC_DEENROLL_CHECK, Z_EC_RATECHANGE, Z_EC_INSTPLAN, Z_EC_BUDGETBILL, Z_EC_CAUPDATE, Z_EC_PAYMENT, Z_EC_ASSIST, Z_EC_CORR, Z_EC_BILLGROUP, Z_EC_SERVICEORDER)"],
], widths=[1.4,0.9,4.4], font_size=8.5)
doc.add_heading("4.3 Principal Propagation (OData reads)", level=2)
numbered([
 "Cloud Connector: enable principal propagation for the subaccount; configure the system CA to issue short-lived X.509 user certificates.",
 "S/4HANA: import the Cloud Connector system CA into STRUST (SSL server Standard); map certificate subject to SU01 users via CERTRULE (rule: CN = e-mail → user alias).",
 "Verify: an OData call from CAP with the agent's propagated identity returns data under that named user's authorizations (test user with restricted contract account display should see 403 on others).",
])
doc.add_heading("4.4 Destinations (Destination Service, identical names all tiers)", level=2)
table(["Name","Type / Proxy","Auth","Key Properties"], [
 ["S4_UTILITIES_ODATA","HTTP / OnPremise","PrincipalPropagation","URL=http://s4-odata:44300; sap-client=<nnn>; CloudConnectorLocationId=EC_MAIN; HTML5.DynamicDestination=true"],
 ["S4_UTILITIES_ODATA_TECH","HTTP / OnPremise","BasicAuthentication (comm user EC_TECH_ODATA)","Fallback for jobs/reconciliation where no user context exists"],
 ["S4_UTILITIES_RFC","RFC / OnPremise","BasicAuthentication (comm user EC_TECH_RFC)","jco.client.mshost/…, jco.destination.repository_destination; used by Cloud Integration RFC adapter"],
 ["SCV2_API","HTTP / Internet","OAuth2ClientCredentials","SC V2 API base URL; client from SC V2 communication system"],
 ["EC_SRV (in SC V2/API Mgmt)","HTTP / Internet","OAuth2","Enrollment Service API for SC V2-side consumption"],
], widths=[1.6,1.2,1.6,2.3], font_size=8.5)

# ========== 5 INTEGRATION SUITE ==========
doc.add_heading("5. Integration Suite Configuration", level=1)
numbered([
 "Provision Integration Suite in each subaccount (or central integration subaccount per client standard); activate capabilities: Cloud Integration, API Management.",
 "Create credential artifacts in the Cloud Integration keystore/security material: EC_TECH_RFC (user credentials), EC_TECH_ODATA (user credentials), SCV2_OAUTH (OAuth2 client credentials), EVENT_MESH_OAUTH.",
 "Import and deploy iFlow package 'EC_EnrollmentCenter' with externalized parameters per tier (hosts, clients, queue names):",
])
table(["iFlow","Trigger / Adapter in-out","Function"], [
 ["IF_EC_Context_Read (IF-01)","HTTPS ← CAP; OData → S4_UTILITIES_ODATA","Parallel multi-entity context read, one aggregate response"],
 ["IF_EC_Eligibility (IF-02)","HTTPS ← CAP; RFC → Z_EC_ELIGIBILITY_PRECHECK","Eligibility batch call; maps RES_EC_ELIG to API JSON"],
 ["IF_EC_Fulfillment_Dispatch (IF-03)","AMQP ← utilco/ec/fulfillment/dispatch; RFC/OData → recipe wrapper","Recipe router: per recipeType to Z_EC_* wrapper; retry 3x exponential; DLQ on exhaust"],
 ["IF_EC_Deenroll (IF-04)","HTTPS ← CAP; RFC → Z_EC_DEENROLL_CHECK + wrapper","Exit validation + reversal execution"],
 ["IF_EC_Status_Events (IF-05)","AMQP ← utilco/s4/enrollment/#; HTTPS → CAP inbox","Confirmation events to EventInbox"],
 ["IF_EC_SCV2_Timeline (IF-06)","HTTPS ← CAP; REST → SCV2_API","Case timeline write-back"],
 ["IF_EC_Exception (IF-08)","DLQ consumer; → SBPA API","Opens exception workflow instances"],
 ["IF_EC_MoveIn_Events (IF-09)","AMQP ← utilco/s4/movein/#; HTTPS → CAP inbox","Move-in completed/cancelled processing"],
], widths=[1.9,2.2,2.4], font_size=8.5)
p("API Management: create API proxy 'enrollment-v1' in front of the CAP service; product 'EC-Internal' "
  "(SC V2 + UI consumption) with spike arrest and JWT validation policies; publish the OpenAPI spec.")

# ========== 6 EVENT MESH ==========
doc.add_heading("6. Event Mesh Configuration", level=1)
table(["Object","Value"], [
 ["Message client","ec-<tier> (namespace utilco/ec/<tier>)"],
 ["Queues","EC_FULFILLMENT_DISPATCH (sub: utilco/ec/fulfillment/dispatch); EC_S4_ENROLLMENT (sub: utilco/s4/enrollment/#); EC_S4_MOVEIN (sub: utilco/s4/movein/#); EC_S4_BPEM (sub: utilco/s4/bpem/#)"],
 ["Queue settings","Max delivery 5 → dead-letter queue EC_DLQ; TTL 7 days; access via OAuth service keys"],
 ["Consumers","Cloud Integration AMQP sender channels (IF-03/05/09); CAP webhook alternative acceptable for low volume"],
 ["Producers","S/4HANA event enablement (Section 9.6) and CAP (dispatch topic)"],
], widths=[1.9,4.6], font_size=9)

# ========== 7 SBPA / HANA ==========
doc.add_heading("7. Build Process Automation and HANA Cloud", level=1)
bullets([
 ("SBPA","import process package EC_Workflows: EC_SupervisorOverride (form: justification, message context; recipient: role EC_SUPERVISOR) and EC_ExceptionReview (retry / correct / cancel actions calling CAP callbacks). Configure destination to CAP API; assign SBPA business roles to IAS groups."),
 ("HANA Cloud","provision per tier; map to CF space; HDI container created by deploy. Enable automated backups (PRD: point-in-time recovery). No direct SQL access for business users."),
 ("CI/CD","BTP CI/CD service: pipeline 'ec-build' on the application GitHub repo (webhook on main); stages build (mbt) → deploy DEV → automated API tests; release stage uploads MTA to Transport Management for QA/PRD import."),
])

# ========== 8 SCV2 ==========
doc.add_heading("8. SAP Service Cloud V2 Configuration", level=1)
numbered([
 "Communication system 'BTP Enrollment Center': OAuth client for the SC V2 APIs used by IF-06 (case service scopes only).",
 "External application registration: Enrollment Center UI URL (approuter route) as mashup; add as tab/work center entry per business role (Utilities Agent) — standalone access point.",
 "Move-In guided flow: embed the same URL as in-screen step (context parameters: businessPartnerId, moveInId, premiseId, requestedStart, mode=MOVE_IN) — validated in the Discover-phase spike; fallback side pane per Solution Design.",
 "Context token: configure signed context handover (agent, BP, case) per SC V2 external app integration; token audience = Enrollment Center XSUAA client.",
 "Autoflow (optional): on timeline entry type 'EC_STATUS', create follow-up task for rejected enrollments.",
])

# ========== 9 S/4 CONFIG ==========
doc.add_page_break()
doc.add_heading("9. S/4HANA Utilities — Required Configuration", level=1)
doc.add_heading("9.1 Gateway / OData", level=2)
numbered([
 "Activate SAP Gateway (SICF nodes /sap/opu/odata) and register services: API_BUSINESS_PARTNER, CDS-based utilities services used from TDD Section 4, Z_EC_CONTEXT_SRV, Z_EC_ENROLLMENT_SRV (/IWFND/MAINT_SERVICE, system alias LOCAL).",
 "Enable OData V4 service groups where CDS service bindings are V4 (publish via /IWFND/V4_ADMIN).",
])
doc.add_heading("9.2 Users and Authorizations", level=2)
table(["User","Type","Role (PFCG) Content"], [
 ["EC_TECH_ODATA","Communication/System","S_SERVICE for registered OData services (reconciliation reads); display auths: BP, CA, utilities master data"],
 ["EC_TECH_RFC","Communication/System","S_RFC for Z_EC_* function group; business auths for wrapped operations: rate change (installation change), FI-CA installment plan create, CA change, service order create — least privilege per wrapper"],
 ["Named agents (propagated)","Dialog/Comm via CERTRULE mapping","Existing utilities agent display roles; no change to their write auths (writes go through EC_TECH_RFC with agent ID logged)"],
], widths=[1.4,1.4,3.7], font_size=8.5)
doc.add_heading("9.3 Security / Trust", level=2)
numbered([
 "STRUST: import Cloud Connector system CA certificate (SSL server Standard PSE).",
 "CERTRULE: certificate mapping rule (subject CN/e-mail → user) for principal propagation.",
 "SM59: no outbound RFC needed for EC; HTTPS outbound to Event Mesh per Section 9.6 option chosen.",
])
doc.add_heading("9.4 BRFplus Prerequisites", level=2)
numbered([
 "Confirm BRFplus active (transaction BRF+ / BRFplus); set client-specific customizing options; storage type S for ZEC_PROGRAM_ELIGIBILITY (transportable).",
 "Transport: workbench request for application/functions/rulesets; enable 'generate on activation'; schedule regeneration job FDT_TRANS after import in QA/PRD.",
 "Authorization: maintainers need S_FDT_* (rule content maintenance restricted to decision tables via authorization profile for business rule owners).",
])
doc.add_heading("9.5 Functional Configuration Touchpoints", level=2)
table(["Area","Configuration","Used By"], [
 ["FI-CA installment plan","Installment plan categories/variants for DPA (interest, min amount, max installments)","Z_EC_INSTPLAN (DPA recipe)"],
 ["Budget billing","Budget billing procedure parameters per division","Z_EC_BUDGETBILL"],
 ["Contract account","Alternative due-date rules for CSDD; payment methods for AutoPay (direct debit mandate settings)","Z_EC_CAUPDATE / Z_EC_PAYMENT"],
 ["Rates/products","Target rate categories & products for DPP/CPP/TOU/PTR/Green/Community Solar variants; operands for surcharges/discounts; allowed-rates matrix (re-used legacy Z_MIMO_* tables)","Z_EC_RATECHANGE"],
 ["Assistance","CA/BP status marks for CARE/FERA/LIHEAP/PIPP; discount operands; protected-customer dunning variant (per CS-02)","Z_EC_ASSIST"],
 ["Correspondence","Bill copy recipient / additional correspondence recipient variants","Z_EC_CORR (TPN, eBill)"],
 ["Joint invoicing","Bill grouping / collective invoicing settings for Summary Billing","Z_EC_BILLGROUP"],
 ["BPEM","Case categories ZEC1 (fulfillment error), ZEC2 (data error); solution processes; auto-close on resolution event","Exception handling"],
 ["Number ranges","Intervals for ZEC_ENROLL_HDR / _HIST","Z_EC_ENROLLMENT_SRV"],
], widths=[1.3,3.4,1.8], font_size=8.5)
doc.add_heading("9.6 Event Enablement", level=2)
p("Choose per installed release (open point 2 of the TDD):")
bullets([
 ("Option A (preferred, ≥2022)","SAP event enablement add-on for on-premise: configure outbound binding to Event Mesh (channel, topic mapping utilco/s4/…); attach to change events for move-in completion, enrollment confirmation (raised inside Z_EC_* wrappers via ABAP event API), BPEM case status."),
 ("Option B (fallback)","Z outbound publisher: wrapper FMs and BPEM exit call a small Z class posting to Event Mesh REST (destination-based) in update task — same topics, at-least-once with EventInbox dedup on BTP."),
])

# ========== 10 S/4 DEVELOPMENTS ==========
doc.add_page_break()
doc.add_heading("10. S/4HANA Utilities — Required Developments (WRICEF)", level=1)
p("Development inventory with specification detail. All objects in package ZEC_ENROLLMENT, prefix Z_EC_/ZEC_. "
  "Every RFC validates its own authorization object check and writes an application log (SLG1 object ZEC).")
doc.add_heading("10.1 Interfaces / Services", level=2)
table(["ID","Object","Type","Specification Summary","Cplx"], [
 ["D-01","Z_EC_ELIGIBILITY_PRECHECK","RFC FM","IMPORT: BP, tab of CA, mode (STANDALONE/PRE_ACTIVE), program filter (opt), move-in ID (opt). Builds CTX_EC_ELIG via parallel aRFC context reads (rate cat from EANLH, device AMI via I_UtilsAdvncdMeteringSystem, dunning/returned-payments from FI-CA, program history from ZEC_ENROLL_HIST). Calls BRFplus F_EC_ELIGIBILITY_ALL once; EXPORT: RES_EC_ELIG table + return. Perf target <1.5 s for 2 CAs × all programs.","H"],
 ["D-02","Z_EC_DEENROLL_CHECK","RFC FM","IMPORT: enrollment key. Evaluates F_EC_DEENROLL_CHECK (min participation via ZEC_ENROLL_HDR dates, open installment balance via FI-CA, pending service orders); EXPORT: blocked flag + messages.","M"],
 ["D-03","Z_EC_CONTEXT_SRV","OData (V4 service binding)","Consolidated read projection over TDD Section 4 CDS views + FI-CA facts (balance, dunning level, locks) exposed as one $batch-friendly service; read-only; principal propagation.","M"],
 ["D-04","Z_EC_ENROLLMENT_SRV","OData","CRUD on ZEC_ENROLL_HDR/_HIST with ETag, change documents, authority check ZEC_ENRL.","M"],
], widths=[0.5,1.6,0.9,3.1,0.4], font_size=8)
doc.add_heading("10.2 Fulfillment Wrappers (RFC, called by IF-03/IF-04)", level=2)
table(["ID","Object","Wraps / Core Logic","Programs","Cplx"], [
 ["D-05","Z_EC_RATECHANGE","Installation rate category change (EANLH/TARIFTYP) with effective date; validates against allowed-rates matrix (re-use Z_MIMO_MI/MO_RES_ALLOWED_RATES, Z_GET_NEW_MI_RATE_CAT); operand add for Green/CSOL surcharge; commit + event raise","DPP, PTR, CPP, TOU, GREEN, CSOL, INTR","H"],
 ["D-06","Z_EC_INSTPLAN","FI-CA installment plan create/deactivate from open items selection; down-payment handling; dunning lock while active","DPA","H"],
 ["D-07","Z_EC_BUDGETBILL","Budget billing plan create/end per procedure","BB","M"],
 ["D-08","Z_EC_CAUPDATE","Contract account changes: CSDD due-date rule, eBill flag, prepay settings","CSDD, EBILL, PREPAY","M"],
 ["D-09","Z_EC_PAYMENT","Payment method/mandate set (AutoPay); due-date deferral (Payment Extension); AMP forgiveness schedule postings","AUTOPAY, PEXT, AMP","H"],
 ["D-10","Z_EC_ASSIST","BP/CA assistance status marks + discount operand + protected dunning variant switch","CARE, FERA, LIHEAP, PIPP, MED","M"],
 ["D-11","Z_EC_CORR","Additional correspondence recipient / bill copy control","TPN","L"],
 ["D-12","Z_EC_BILLGROUP","Joint invoicing / bill grouping assignment","SUMBILL","M"],
 ["D-13","Z_EC_SERVICEORDER","Service order create for device install/exchange with code groups","THERM, PPS, AMI exchange","M"],
 ["D-14","Z_EC_MOVEIN_SRV","Move-in create/read wrapper for the SC V2 guided flow (only if standard SC V2 utilities move-in integration is not licensed)","Move-in mode","H"],
], widths=[0.5,1.3,2.9,1.3,0.4], font_size=8)
doc.add_heading("10.3 Dictionary, Rules and Events", level=2)
table(["ID","Object","Type","Specification Summary","Cplx"], [
 ["D-15","ZEC_ENROLL_HDR / ZEC_ENROLL_HIST","Tables + CDS","Participation header (BP, CA, program, status, request/actual dates, channel, reference keys) and history; change documents; archiving object","M"],
 ["D-16","ZEC_PROGRAM_ELIGIBILITY","BRFplus app","Per TDD Section 6: functions F_EC_ELIGIBILITY_ALL / F_EC_DEENROLL_CHECK, 28 rulesets, decision tables DT_*, context/result structures; legacy RS_DPP_INT_VALIDATION migrated (Dynamic Peak Pricing)","H"],
 ["D-17","ZEC_ELIG","Message class","0xx hard stops, 3xx warnings, 9xx not-yet-eligible; all UI/audit messages","L"],
 ["D-18","Event raising","ABAP (in wrappers + BPEM exit)","Raise utilco/s4/enrollment/confirmed, movein completed/cancelled, bpem casechanged per Section 9.6 option","M"],
 ["D-19","ZEC_RECON_EXTRACT","Report/OData","Daily reconciliation extract: participation + in-flight fulfillments for BTP comparison job","L"],
 ["D-20","Authorization objects ZEC_ENRL / ZEC_OVRD","Auth objects","Activity-level control on enrollment writes and override acceptance","L"],
], widths=[0.5,1.6,0.9,3.1,0.4], font_size=8)

# ========== 11 SEQUENCE ==========
doc.add_page_break()
doc.add_heading("11. Deployment Sequence (per tier)", level=1)
numbered([
 "BTP: create subaccount, enable CF, assign entitlements (Section 2); establish IAS trust and role collections (Section 3).",
 "Install/connect Cloud Connector; expose virtual hosts and resources (Section 4.1–4.2).",
 "S/4HANA: create packages, users, roles (9.2); import development transports D-01…D-20 (Section 10); activate OData services (9.1); STRUST/CERTRULE (9.3); functional configuration (9.5); BRFplus transport + generation (9.4).",
 "BTP: create destinations (4.4) and test connectivity (destination check + smoke OData/RFC ping).",
 "Provision HANA Cloud; deploy MTA (CAP + UI) via CI/CD (Sections 2.3, 7).",
 "Provision Event Mesh; create queues/subscriptions (6); configure S/4HANA event enablement (9.6); verify event round trip.",
 "Deploy Integration Suite package and credentials (5); wire CAP ↔ iFlow endpoints (externalized params).",
 "Import SBPA package; map roles (7).",
 "SC V2: communication system, mashup/work center, move-in step embedding, context token (8).",
 "Execute validation checklist (Section 12); obtain sign-off before next tier.",
])

# ========== 12 VALIDATION ==========
doc.add_heading("12. Post-Deployment Validation Checklist", level=1)
table(["#","Check","Expected"], [
 ["V-01","Open Enrollment Center from SC V2 work center as test agent","UI loads <3 s with context; no re-login; four panels populated"],
 ["V-02","Eligibility for seeded test customer (2 CAs)","One IF-02 call; messages with correct severities; log entry in SLG1 ZEC"],
 ["V-03","Enroll CSDD (simplest recipe)","Status Enrollment Request → Enrollment Pending → Enrolled; CA due-date rule visible in S/4HANA; timeline entry on SC V2 case"],
 ["V-04","Enroll DPP (rate change recipe)","Product/rate change document created effective next cycle; event confirmed"],
 ["V-05","DPA on account with arrears","Installment plan created; dunning lock set; eligibility now hard-stops second DPA"],
 ["V-06","Warning override path (CSDD with returned payments)","SBPA task to supervisor; approval recorded; enrollment proceeds; OverrideRecord + audit written"],
 ["V-07","Move-in embedded mode: park + activate","Request parked; move-in completed event releases it; re-validation logged; cancellation cascade works on second test"],
 ["V-08","De-enrollment blocked by exit rule","422 with EC_EXIT message; no backend change"],
 ["V-09","Kill S/4HANA connection during fulfillment","Retry then DLQ; status Error – In Review; SBPA exception task; recovery by retry action"],
 ["V-10","Principal propagation authorization","Restricted test agent receives 403 on out-of-scope contract account read"],
 ["V-11","Reconciliation job","Deliberate missed event detected and flagged within one run"],
], widths=[0.6,2.8,3.1], font_size=8.5)

# ========== APPENDIX ==========
doc.add_heading("Appendix A — xs-security.json (skeleton)", level=1)
mono('{ "xsappname": "ec-enrollment", "tenant-mode": "dedicated",')
mono('  "scopes": [ {"name": "$XSAPPNAME.Agent"}, {"name": "$XSAPPNAME.Supervisor"}, {"name": "$XSAPPNAME.Admin"} ],')
mono('  "role-templates": [')
mono('    {"name": "ECAgent", "scope-references": ["$XSAPPNAME.Agent"]},')
mono('    {"name": "ECSupervisor", "scope-references": ["$XSAPPNAME.Agent","$XSAPPNAME.Supervisor"]},')
mono('    {"name": "ECAdmin", "scope-references": ["$XSAPPNAME.Agent","$XSAPPNAME.Supervisor","$XSAPPNAME.Admin"]} ],')
mono('  "oauth2-configuration": { "redirect-uris": ["https://*.hana.ondemand.com/**"] } }')
doc.add_heading("Appendix B — Naming Conventions", level=1)
bullets([
 "S/4HANA: package ZEC_ENROLLMENT; objects Z_EC_* (executables) / ZEC_* (dictionary, message class, BRFplus, auth objects).",
 "BTP: MTA ec-enrollment; destinations S4_UTILITIES_*, SCV2_API; role collections EnrollmentCenter_*.",
 "Event topics: utilco/s4/* (backend-produced), utilco/ec/* (platform-produced).",
 "iFlows: IF_EC_<Function> matching IF-nn IDs from the Solution Design.",
])

out = r"C:\Users\jnamm\OneDrive\Desktop\Service Cloud for Utilities\Enrollment Center\Enrollment Center - BTP Deployment and Configuration Guide v1.0.docx"
doc.save(out)
print("Saved:", out)
